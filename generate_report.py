"""
generate_report.py - Gerador de Relatórios Executivos PROFISSIONAL
Versão Enterprise com gráficos de criticidade por arquivo
MELHORIAS:
  - Donut chart de distribuição de severidade
  - Gráfico de barras: score de criticidade por arquivo
  - Gráfico de barras: volume exposto por arquivo
  - Mapa de calor multidimensional de risco
  - Gauge individual (velocímetro) para cada arquivo vulnerável
  - Suporte completo AWS S3 / GCS / Azure Blob Storage
"""

# ── Matplotlib (backend não-interativo ANTES de qualquer import plt) ──
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

# ── ReportLab ─────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
    PageBreak, KeepTogether, Image as RLImage
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT

# ── Stdlib ────────────────────────────────────────────────────────────
from datetime import datetime
from pathlib import Path
from collections import Counter
import hashlib, json, io, tempfile, os

# ── python-docx ───────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.shared import RGBColor, Pt, Inches, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ══════════════════════════════════════════════════════════════════════
# CHART ENGINE — gera PNGs em memória (BytesIO) ou arquivo temporário
# ══════════════════════════════════════════════════════════════════════
class ChartEngine:
    """Gera todos os gráficos como imagens PNG prontas para embutir."""

    PALETTE = {
        'primary':   '#1e3a8a',
        'secondary': '#3b82f6',
        'critical':  '#dc2626',
        'high':      '#ea580c',
        'medium':    '#ca8a04',
        'low':       '#16a34a',
        'bg':        '#f8fafc',
        'gray':      '#6b7280',
        'white':     '#ffffff',
    }

    SEV_COLOR = {
        'CRITICAL': '#dc2626',
        'HIGH':     '#ea580c',
        'MEDIUM':   '#ca8a04',
        'LOW':      '#16a34a',
    }

    # Score base por severidade (0-100) — cada severidade fica dentro da sua zona
    # LOW: 0-25 | MEDIUM: 25-50 | HIGH: 50-75 | CRITICAL: 75-100
    SEV_BASE = {'CRITICAL': 87, 'HIGH': 62, 'MEDIUM': 37, 'LOW': 12}

    def __init__(self, tmp_dir: Path):
        self.tmp = tmp_dir
        self.tmp.mkdir(exist_ok=True)

    def _save(self, fig, name: str) -> str:
        """Salva figura como PNG e retorna o caminho."""
        path = str(self.tmp / f'{name}.png')
        fig.savefig(path, dpi=200, bbox_inches='tight',
                    facecolor='white', edgecolor='none')
        plt.close(fig)
        return path

    def _file_score(self, f: dict, max_size: float) -> float:
        """Calcula score de criticidade de um arquivo (0-100)."""
        base  = self.SEV_BASE.get(f.get('severity', 'LOW'), 18)
        size  = float(f.get('size', 0))
        bonus = (size / max_size * 12) if max_size > 0 else 0
        return min(100.0, base + bonus)

    # ── 1. Donut — distribuição de severidade ─────────────────────────
    def severity_donut(self, severity_dist: dict, total: int) -> str:
        P = self.PALETTE
        labels  = ['CRÍTICO', 'ALTO', 'MÉDIO', 'BAIXO']
        keys    = ['critical', 'high', 'medium', 'low']
        pal     = [P['critical'], P['high'], P['medium'], P['low']]
        values  = [severity_dist.get(k, 0) for k in keys]

        nz_vals = [v for v in values if v > 0]
        nz_cols = [c for c, v in zip(pal, values) if v > 0]
        nz_lbls = [f'{l}: {v:,}' for l, v in zip(labels, values) if v > 0]

        fig, ax = plt.subplots(figsize=(6.5, 5.2), facecolor='white')

        if sum(nz_vals) > 0:
            wedges, _, autotexts = ax.pie(
                nz_vals, colors=nz_cols,
                autopct=lambda p: f'{p:.1f}%' if p > 2 else '',
                startangle=90, pctdistance=0.78,
                wedgeprops={'linewidth': 2.5, 'edgecolor': 'white',
                            'antialiased': True}
            )
            for at in autotexts:
                at.set_fontsize(10); at.set_fontweight('bold'); at.set_color('white')
            ax.legend(wedges, nz_lbls, loc='lower center',
                      bbox_to_anchor=(0.5, -0.08), ncol=2,
                      fontsize=9.5, frameon=False,
                      handlelength=1.2, handleheight=1.0)

        # Buraco central com sombra
        shadow = plt.Circle((0.03, -0.03), 0.56, color='#e5e7eb', zorder=1)
        circle = plt.Circle((0, 0), 0.56, color='white', zorder=2)
        ax.add_patch(shadow)
        ax.add_patch(circle)

        dominant_sev = labels[values.index(max(values))] if any(v > 0 for v in values) else 'BAIXO'
        dom_col      = pal[values.index(max(values))] if any(v > 0 for v in values) else P['low']
        ax.text(0,  0.10, f'{total:,}',   ha='center', va='center', zorder=3,
                fontsize=22, fontweight='bold', color=P['primary'])
        ax.text(0, -0.14, 'arquivos',    ha='center', va='center', zorder=3,
                fontsize=9.5, color=P['gray'])
        ax.text(0, -0.34, dominant_sev, ha='center', va='center', zorder=3,
                fontsize=8.5, fontweight='bold', color=dom_col)

        ax.set_title('Distribuição de Severidade', fontsize=13,
                     fontweight='bold', color=P['primary'], pad=14)
        ax.axis('equal')
        fig.tight_layout(pad=1.2)
        return self._save(fig, 'donut_severity')

    # ── 2. Barras — score de criticidade por arquivo ──────────────────
    def risk_per_file(self, files: list, title: str = '', fname: str = 'risk_files') -> str:
        if not files:
            return None
        P = self.PALETTE
        max_size = max(float(f.get('size', 0)) for f in files) or 1

        names  = [f.get('key', f.get('name', ''))[-30:] for f in files]
        scores = [self._file_score(f, max_size) for f in files]
        cols   = [self.SEV_COLOR.get(f.get('severity', 'LOW'), P['low']) for f in files]

        n = len(files)
        fig, ax = plt.subplots(figsize=(9.5, max(4.0, n * 0.50 + 1.4)), facecolor='white')
        ax.set_facecolor('white')

        y = np.arange(n)
        # Zonas de fundo sutis primeiro
        for x0, x1, c in [(0,25,P['low']),(25,50,P['medium']),(50,75,P['high']),(75,100,P['critical'])]:
            ax.axvspan(x0, x1, alpha=0.055, color=c, zorder=0)

        bars = ax.barh(y, scores, color=cols, height=0.62,
                       edgecolor='none', linewidth=0, alpha=0.88, zorder=2)
        # Linha de borda esquerda colorida
        for bar, col in zip(bars, cols):
            ax.barh(bar.get_y(), 0.5, height=bar.get_height(),
                    color=col, alpha=1.0, zorder=3, left=0)

        ax.grid(axis='x', color='#f3f4f6', linewidth=1.2, zorder=1)
        ax.set_axisbelow(True)

        for bar, score, f in zip(bars, scores, files):
            sev = f.get('severity', 'LOW')
            ax.text(min(bar.get_width() + 1.2, 98),
                    bar.get_y() + bar.get_height() / 2,
                    f'{sev}  {score:.0f} pts',
                    va='center', ha='left', fontsize=8,
                    fontweight='bold', color=self.SEV_COLOR.get(sev, P['low']))

        for x, lbl, c in [(12.5,'BAIXO',P['low']),(37.5,'MÉDIO',P['medium']),
                           (62.5,'ALTO',P['high']),(87.5,'CRÍTICO',P['critical'])]:
            ax.text(x, n - 0.05, lbl, ha='center', fontsize=7.5,
                    color=c, fontweight='bold', style='italic', va='bottom')

        ax.set_yticks(y)
        ax.set_yticklabels(names, fontsize=8.5, color='#374151')
        ax.set_xlim(0, 105)
        ax.set_xlabel('Score de Criticidade', fontsize=9, color=P['gray'])
        ax.set_title(f'Score de Criticidade por Arquivo — Top 20',
                     fontsize=12, fontweight='bold', color=P['primary'], pad=12)
        for sp in ['top', 'right', 'left']:
            ax.spines[sp].set_visible(False)
        ax.spines['bottom'].set_color('#e5e7eb')
        ax.tick_params(left=False)
        ax.invert_yaxis()
        fig.tight_layout(pad=1.1)
        return self._save(fig, fname)

    # ── 3. Barras — tamanho por arquivo ──────────────────────────────
    def size_per_file(self, files: list) -> str:
        if not files:
            return None
        P = self.PALETTE

        def fmt_size(b):
            b = float(b)
            for u in ['B', 'KB', 'MB', 'GB']:
                if b < 1024: return b, u
                b /= 1024
            return b, 'TB'

        names = [f.get('key', f.get('name', ''))[-30:] for f in files]
        sizes_raw = [float(f.get('size', 0)) for f in files]
        max_s = max(sizes_raw) or 1
        sizes_mb = [s / (1024*1024) for s in sizes_raw]
        cols = [self.SEV_COLOR.get(f.get('severity', 'LOW'), P['low']) for f in files]

        n = len(files)
        fig, ax = plt.subplots(figsize=(9.5, max(3.5, n * 0.44 + 1.2)), facecolor='white')
        ax.set_facecolor('#f9fafb')

        y = np.arange(n)
        bars = ax.barh(y, sizes_mb, color=cols, height=0.58,
                       edgecolor='white', linewidth=0.8, alpha=0.88)
        ax.grid(axis='x', color='white', linewidth=1.4)
        ax.set_axisbelow(True)

        for bar, raw in zip(bars, sizes_raw):
            val, unit = fmt_size(raw)
            ax.text(bar.get_width() + max(sizes_mb) * 0.01,
                    bar.get_y() + bar.get_height() / 2,
                    f'{val:.2f} {unit}', va='center', ha='left',
                    fontsize=8.5, color='#374151')

        ax.set_yticks(y)
        ax.set_yticklabels(names, fontsize=8.5, color='#374151')
        ax.set_xlabel('Tamanho (MB)', fontsize=9, color=P['gray'])
        ax.set_title('Volume de Dados Expostos por Arquivo',
                     fontsize=12, fontweight='bold', color=P['primary'], pad=10)
        for sp in ['top', 'right']:
            ax.spines[sp].set_visible(False)
        ax.invert_yaxis()
        fig.tight_layout()
        return self._save(fig, 'size_per_file')

    # ── 4. Mapa de calor multidimensional ────────────────────────────
    def risk_heatmap(self, files: list) -> str:
        if not files:
            return None
        P   = self.PALETTE
        n   = min(len(files), 15)
        sel = files[:n]
        max_size = max(float(f.get('size', 0)) for f in sel) or 1

        names = [f.get('key', f.get('name', ''))[-22:] for f in sel]
        dims  = ['Exposição\nPública', 'Volume\nDados', 'Tipo\nArquivo',
                 'Risco\nPrivacidade', 'Score\nFinal']

        def row_scores(f):
            sev  = f.get('severity', 'LOW')
            size = float(f.get('size', 0))
            ext  = Path(f.get('key', f.get('name', ''))).suffix.lower()
            # Exposição: crítico=95, alto=75, médio=50, baixo=30
            exp  = {'CRITICAL':95,'HIGH':75,'MEDIUM':50,'LOW':30}.get(sev, 30)
            vol  = min(100, size / max_size * 100)
            typ  = 65 if ext in ['.sql','.db','.env','.key','.pem'] else \
                   55 if ext in ['.zip','.tar','.gz','.bak'] else \
                   50 if ext in ['.js','.py','.php'] else 40
            priv = 70 if any(x in f.get('key', f.get('name','')).lower()
                             for x in ['user','customer','person','cpf','password','secret']) else 45
            final = min(100, self._file_score(f, max_size))
            return [exp, vol, typ, priv, final]

        matrix = np.array([row_scores(f) for f in sel])

        fig, ax = plt.subplots(figsize=(10.5, max(4, n * 0.48 + 1.8)), facecolor='white')
        im = ax.imshow(matrix, cmap=plt.cm.RdYlGn_r, vmin=0, vmax=100, aspect='auto')

        for i in range(n):
            for j in range(len(dims)):
                val = matrix[i, j]
                tc  = 'white' if val > 65 or val < 20 else '#222222'
                ax.text(j, i, f'{val:.0f}', ha='center', va='center',
                        fontsize=8.5, fontweight='bold', color=tc)

        ax.set_xticks(range(len(dims)))
        ax.set_xticklabels(dims, fontsize=9, fontweight='bold', color=P['primary'])
        ax.set_yticks(range(n))
        ax.set_yticklabels(names, fontsize=8.5, color='#374151')
        ax.set_title('Mapa de Calor de Risco por Arquivo',
                     fontsize=12, fontweight='bold', color=P['primary'], pad=12)

        cbar = fig.colorbar(im, ax=ax, fraction=0.025, pad=0.02)
        cbar.set_label('Nível de Risco (0–100)', fontsize=8, color=P['gray'])
        cbar.ax.tick_params(labelsize=7)
        fig.tight_layout()
        return self._save(fig, 'heatmap')

    # ── 5. Pizza — distribuição por tipo de arquivo ───────────────────
    def type_pie(self, files: list) -> str:
        if not files:
            return None
        P = self.PALETTE
        counts = Counter(
            Path(f.get('key', f.get('name', ''))).suffix.lower() or '(sem ext)'
            for f in files
        )
        # Agrupar raridades
        top = counts.most_common(6)
        others = sum(v for _, v in counts.most_common()[6:])
        if others:
            top.append(('outros', others))

        labels = [k for k, _ in top]
        values = [v for _, v in top]
        pal = [P['secondary'], P['primary'], P['critical'],
               P['medium'], P['low'], P['high'], P['gray']]

        fig, ax = plt.subplots(figsize=(6, 4.5), facecolor='white')
        wedges, texts, autotexts = ax.pie(
            values, labels=labels, colors=pal[:len(labels)],
            autopct='%1.1f%%', startangle=90,
            wedgeprops={'linewidth': 2, 'edgecolor': 'white'},
            pctdistance=0.78
        )
        for at in autotexts:
            at.set_fontsize(10); at.set_fontweight('bold'); at.set_color('white')
        for t in texts:
            t.set_fontsize(9); t.set_color('#374151')

        ax.set_title('Distribuição por Tipo de Arquivo',
                     fontsize=12, fontweight='bold', color=P['primary'], pad=12)
        fig.tight_layout()
        return self._save(fig, 'type_pie')

    # ── 6. Gauge individual por arquivo ──────────────────────────────
    def individual_gauge(self, f: dict, idx: int, max_size: float) -> str:
        score = self._file_score(f, max_size)
        sev   = f.get('severity', 'LOW')
        color = self.SEV_COLOR.get(sev, self.PALETTE['low'])
        name  = f.get('key', f.get('name', ''))

        fig = plt.figure(figsize=(5.5, 3.4), facecolor='white')
        ax  = fig.add_subplot(111, projection='polar')
        ax.set_facecolor('white')

        # Zonas coloridas em semicírculo (π → 0)
        zones = [
            (np.pi,       np.pi*0.75, self.PALETTE['low']),
            (np.pi*0.75,  np.pi*0.50, self.PALETTE['medium']),
            (np.pi*0.50,  np.pi*0.25, self.PALETTE['high']),
            (np.pi*0.25,  0,          self.PALETTE['critical']),
        ]
        for t0, t1, zcol in zones:
            th = np.linspace(t0, t1, 60)
            ax.fill_between(th, 0.62, 0.98, alpha=0.20, color=zcol)
            ax.plot(th, np.full(60, 0.98), color=zcol, linewidth=4.5,
                    solid_capstyle='butt')

        # Agulha
        needle = np.pi - (score / 100.0) * np.pi
        ax.annotate('', xy=(needle, 0.82), xytext=(0, 0),
                    arrowprops=dict(arrowstyle='->', color=color,
                                   lw=3.5, mutation_scale=20))

        ax.set_ylim(0, 1.1)
        ax.axis('off')

        # Score e severidade no centro
        ax.text(0,  0.10, f'{score:.0f}', ha='center', va='center',
                fontsize=22, fontweight='bold', color=color,
                transform=ax.transData)
        ax.text(0, -0.20, sev, ha='center', va='center',
                fontsize=10, fontweight='bold', color=color,
                transform=ax.transData)

        # Rótulos das zonas
        for ang, lbl, c in [
            (np.pi*0.875, 'LOW',  self.PALETTE['low']),
            (np.pi*0.625, 'MED',  self.PALETTE['medium']),
            (np.pi*0.375, 'HIGH', self.PALETTE['high']),
            (np.pi*0.125, 'CRIT', self.PALETTE['critical']),
        ]:
            ax.text(ang, 1.12, lbl, ha='center', va='center',
                    fontsize=7, color=c, fontweight='bold')

        short = (name[-22:] if len(name) > 22 else name)
        ax.set_title(short, fontsize=9, fontweight='bold',
                     color='#374151', pad=2, y=0.08)
        fig.tight_layout(pad=0.4)
        return self._save(fig, f'gauge_{idx:03d}')

    # ── 7. Gauge global do ambiente ───────────────────────────────────
    def environment_gauge(self, score: float, risk_level: str) -> str:
        P     = self.PALETTE
        color = self.SEV_COLOR.get(risk_level.upper(), P['medium'])
        fig   = plt.figure(figsize=(5.5, 3.8), facecolor='white')
        ax    = fig.add_axes([0.05, 0.15, 0.90, 0.80], projection='polar')
        ax.set_facecolor('white')

        zones = [
            (np.pi,      np.pi*0.75, P['low']),
            (np.pi*0.75, np.pi*0.50, P['medium']),
            (np.pi*0.50, np.pi*0.25, P['high']),
            (np.pi*0.25, 0,          P['critical']),
        ]
        for t0, t1, zcol in zones:
            th = np.linspace(t0, t1, 80)
            ax.fill_between(th, 0.55, 1.0, alpha=0.18, color=zcol)
            ax.plot(th, np.full(80, 1.0), color=zcol, linewidth=7, solid_capstyle='butt')

        needle = np.pi - (min(score, 100) / 100.0) * np.pi
        ax.annotate('', xy=(needle, 0.82), xytext=(0, 0),
                    arrowprops=dict(arrowstyle='->', color=color, lw=4, mutation_scale=22))

        ax.set_ylim(0, 1.15)
        ax.set_rticks([])          # remove radial tick labels (source of "1.00" bug)
        ax.set_thetagrids([])      # remove angular grid lines
        ax.axis('off')

        for ang, lbl, c in [
            (np.pi*0.875, 'LOW',  P['low']),
            (np.pi*0.625, 'MED',  P['medium']),
            (np.pi*0.375, 'HIGH', P['high']),
            (np.pi*0.125, 'CRIT', P['critical']),
        ]:
            ax.text(ang, 1.18, lbl, ha='center', va='center',
                    fontsize=8, color=c, fontweight='bold')

        # Use fig.text() for center labels — avoids polar coordinate issues
        fig.text(0.50, 0.46, f'{score:.0f}',      ha='center', va='center',
                 fontsize=28, fontweight='bold', color=color)
        fig.text(0.50, 0.33, 'RISK SCORE',         ha='center', va='center',
                 fontsize=9,  fontweight='bold', color=P['gray'])
        fig.text(0.50, 0.22, risk_level.upper(),   ha='center', va='center',
                 fontsize=11, fontweight='bold', color=color)
        fig.text(0.50, 0.06, 'Risco Global do Ambiente', ha='center', va='center',
                 fontsize=10, fontweight='bold', color=P['primary'])
        return self._save(fig, 'env_gauge')

    # ── 8. Barras de categoria ────────────────────────────────────────
    def category_bars(self, files: list) -> str:
        if not files:
            return None
        P = self.PALETTE
        from pathlib import Path as _P
        from collections import Counter

        CAT_RULES = {
            'Credenciais':    lambda k: any(x in k for x in ['password','credential','secret','token','api_key']),
            'Chaves Crypto':  lambda k: any(x in k for x in ['.pem','.key','.crt','.p12','.pfx','rsa']),
            'DB / Backups':   lambda k: any(x in k for x in ['.sql','.db','.sqlite','backup','.bak','.dump','.gz','.zip','.tar']),
            'Código':         lambda k: any(x in k for x in ['.py','.js','.java','.php','.rb','.go']),
            'Configs':        lambda k: any(x in k for x in ['.env','.yaml','.yml','.conf','.ini','config']),
            'Outros':         lambda k: True,
        }
        SEV_PAL = {'CRITICAL': P['critical'], 'HIGH': P['high'], 'MEDIUM': P['medium'], 'LOW': P['low']}

        cats   = list(CAT_RULES.keys())
        counts = {c: {s: 0 for s in ['CRITICAL','HIGH','MEDIUM','LOW']} for c in cats}
        for f in files:
            name = f.get('key', f.get('name', '')).lower()
            sev  = f.get('severity', 'LOW').upper()
            for cat, rule in CAT_RULES.items():
                if rule(name):
                    counts[cat][sev] += 1
                    break

        cats_with_data = [c for c in cats if sum(counts[c].values()) > 0]
        if not cats_with_data:
            return None

        totals = [sum(counts[c].values()) for c in cats_with_data]
        idx    = np.arange(len(cats_with_data))
        n_c    = len(cats_with_data)
        fig, ax = plt.subplots(figsize=(9, max(1.8, n_c*0.85+0.7)), facecolor='white')
        ax.set_facecolor('white')

        max_total   = max(totals)
        grand_total = sum(totals)

        # Só inclui na legenda as severidades que têm dados
        sev_present = [s for s in ['CRITICAL','HIGH','MEDIUM','LOW']
                       if any(counts[c][s] > 0 for c in cats_with_data)]

        left = np.zeros(len(cats_with_data))
        legend_handles = []
        for sev, col in [('CRITICAL',P['critical']),('HIGH',P['high']),('MEDIUM',P['medium']),('LOW',P['low'])]:
            vals = np.array([counts[c][sev] for c in cats_with_data], dtype=float)
            bars = ax.barh(idx, vals, left=left, color=col, height=0.55,
                           edgecolor='white', linewidth=0.8)
            if sev in sev_present:
                legend_handles.append(mpatches.Patch(color=col, label=sev))
            for bar, v in zip(bars, vals):
                if v > 0 and v >= max_total * 0.06:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_y() + bar.get_height()/2,
                            f'{int(v)}', ha='center', va='center', fontsize=8,
                            fontweight='bold', color='white')
            left += vals

        # Rótulo total + percentual fora de cada barra
        for i, tot in enumerate(totals):
            pct = tot / grand_total * 100 if grand_total else 0
            lbl = f'{tot:,}' if n_c == 1 else f'{tot:,}  ({pct:.0f}%)'
            ax.text(tot + max_total * 0.015, i, lbl,
                    va='center', ha='left', fontsize=9, color='#374151', fontweight='bold')

        ax.set_xlim(0, max_total * 1.28)
        ax.set_yticks(idx)
        ax.set_yticklabels(cats_with_data, fontsize=10, color='#374151')
        ax.set_xlabel('Número de Arquivos', fontsize=9, color=P['gray'])
        ax.set_title('Distribuição por Categoria de Risco',
                     fontsize=12, fontweight='bold', color=P['primary'], pad=12)
        # Legenda fora do plot, à direita — sem overlap com barras
        ax.legend(handles=legend_handles, loc='center left', bbox_to_anchor=(1.01, 0.5),
                  fontsize=9, frameon=False)
        ax.grid(axis='x', color='#f3f4f6', linewidth=1.0)
        ax.set_axisbelow(True)
        for sp in ['top','right','left']:
            ax.spines[sp].set_visible(False)
        ax.spines['bottom'].set_color('#e5e7eb')
        ax.tick_params(left=False)
        ax.invert_yaxis()
        fig.tight_layout()
        return self._save(fig, 'category_bars')

    # ── 9. Funil de achados por severidade ───────────────────────────
    def priority_funnel(self, severity_dist: dict, total: int) -> str:
        P = self.PALETTE
        KEYS   = ['critical', 'high', 'medium', 'low']
        LABELS = ['CRÍTICO',  'ALTO', 'MÉDIO',  'BAIXO']
        COLS   = [P['critical'], P['high'], P['medium'], P['low']]
        vals   = [severity_dist.get(k, 0) for k in KEYS]
        active = [(l, v, c) for l, v, c in zip(LABELS, vals, COLS) if v > 0]
        if not active:
            return None
        maxv = max(v for _, v, _ in active)
        n    = len(active)
        fig, ax = plt.subplots(figsize=(7, max(1.2, n * 0.82 + 0.5)), facecolor='white')
        ax.set_facecolor('white')
        for i, (label, val, color) in enumerate(active):
            pct   = val / total * 100 if total > 0 else 0
            bar_w = maxv * 0.15 + (val / maxv) * (maxv * 0.85)
            ax.barh(i, bar_w, color=color, alpha=0.85, height=0.65, left=0, edgecolor='none')
            ax.text(-maxv * 0.025, i, f'{val:,}',
                    va='center', ha='right', fontsize=12, fontweight='bold', color=color)
            ax.text(bar_w * 0.04, i, label,
                    va='center', ha='left', fontsize=9, fontweight='bold', color='white')
            ax.text(bar_w + maxv * 0.025, i, f'{pct:.1f}%',
                    va='center', ha='left', fontsize=9, fontweight='bold', color=color)
        ax.set_xlim(-maxv * 0.22, maxv * 1.28)
        ax.set_ylim(-0.5, n - 0.5)
        ax.axis('off')
        ax.invert_yaxis()
        fig.tight_layout(pad=0.3)
        return self._save(fig, 'priority_funnel')

    # ── IAM: barras de checks por categoria ──────────────────────────
    def iam_check_bars(self, findings: list) -> str:
        if not findings:
            return None
        P = self.PALETTE

        IAM_CATS = {
            'MFA / Autenticação':   lambda f: 'MFA' in f.get('check','') or 'senha' in f.get('check','').lower(),
            'Access Keys':           lambda f: 'Key' in f.get('check','') or 'key' in f.get('check','').lower(),
            'Permissões (Admin)':    lambda f: 'AdministratorAccess' in f.get('check','') or 'wildcard' in f.get('check','').lower() or 'Action=*' in f.get('check',''),
            'Trust / Cross-Account': lambda f: 'Trust' in f.get('check','') or 'Principal' in f.get('check',''),
            'Configuração de Conta': lambda f: f.get('entity_type','') in ('root','account'),
            'Inatividade':           lambda f: 'inativo' in f.get('check','').lower() or 'nunca' in f.get('check','').lower(),
        }
        SEV_PAL = {'CRITICAL': P['critical'], 'HIGH': P['high'], 'MEDIUM': P['medium'], 'LOW': P['low']}

        cats = list(IAM_CATS.keys())
        counts = {c: {s: 0 for s in ['CRITICAL','HIGH','MEDIUM','LOW']} for c in cats}
        other = {'CRITICAL':0,'HIGH':0,'MEDIUM':0,'LOW':0}
        for f in findings:
            sev = f.get('severity','LOW').upper()
            matched = False
            for cat, rule in IAM_CATS.items():
                if rule(f):
                    counts[cat][sev] += 1
                    matched = True
                    break
            if not matched:
                other[sev] += 1
        if sum(other.values()) > 0:
            counts['Outros'] = other
            cats.append('Outros')

        cats_with_data = [c for c in cats if sum(counts.get(c, {}).values()) > 0]
        if not cats_with_data:
            return None

        totals = [sum(counts[c].values()) for c in cats_with_data]
        max_total = max(totals) or 1
        idx = np.arange(len(cats_with_data))

        fig, ax = plt.subplots(figsize=(9, max(2.0, len(cats_with_data) * 0.75 + 0.6)), facecolor='white')
        ax.set_facecolor('white')

        sev_present = [s for s in ['CRITICAL','HIGH','MEDIUM','LOW']
                       if any(counts[c].get(s,0) > 0 for c in cats_with_data)]
        left = np.zeros(len(cats_with_data))
        legend_handles = []
        for sev, col in [('CRITICAL',P['critical']),('HIGH',P['high']),('MEDIUM',P['medium']),('LOW',P['low'])]:
            vals = np.array([counts[c].get(sev,0) for c in cats_with_data], dtype=float)
            bars = ax.barh(idx, vals, left=left, color=col, height=0.55, edgecolor='white', linewidth=0.8)
            if sev in sev_present:
                legend_handles.append(mpatches.Patch(color=col, label=sev))
            for bar, v in zip(bars, vals):
                if v > 0 and v >= max_total * 0.07:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_y() + bar.get_height()/2,
                            f'{int(v)}', ha='center', va='center', fontsize=8, fontweight='bold', color='white')
            left += vals

        for i, tot in enumerate(totals):
            pct = tot / len(findings) * 100 if findings else 0
            ax.text(tot + max_total * 0.015, i, f'{tot}  ({pct:.0f}%)',
                    va='center', ha='left', fontsize=9, color='#374151', fontweight='bold')

        ax.set_xlim(0, max_total * 1.30)
        ax.set_yticks(idx)
        ax.set_yticklabels(cats_with_data, fontsize=10, color='#374151')
        ax.set_xlabel('Número de Findings', fontsize=9, color=P['gray'])
        ax.set_title('Distribuição de Findings IAM por Categoria', fontsize=12, fontweight='bold', color=P['primary'], pad=12)
        ax.legend(handles=legend_handles, loc='center left', bbox_to_anchor=(1.01, 0.5), fontsize=9, frameon=False)
        ax.grid(axis='x', color='#f3f4f6', linewidth=1.0)
        ax.set_axisbelow(True)
        for sp in ['top','right','left']:
            ax.spines[sp].set_visible(False)
        ax.spines['bottom'].set_color('#e5e7eb')
        ax.tick_params(left=False)
        ax.invert_yaxis()
        fig.tight_layout()
        return self._save(fig, 'iam_check_bars')

    # ── Gerar todos os gauges ─────────────────────────────────────────
    def all_gauges(self, files: list) -> list:
        if not files:
            return []
        max_size = max(float(f.get('size', 0)) for f in files) or 1
        paths = []
        for i, f in enumerate(files):
            sev = f.get('severity', 'LOW')
            if sev in ('CRITICAL', 'HIGH', 'MEDIUM', 'LOW'):
                paths.append(self.individual_gauge(f, i, max_size))
                print(f"  📊 Gauge {i+1}/{len(files)}: {f.get('key', f.get('name',''))[-35:]}")
        return paths


# ══════════════════════════════════════════════════════════════════════
# MAIN GENERATOR CLASS
# ══════════════════════════════════════════════════════════════════════
class ProfessionalReportGenerator:
    """Gerador de relatórios executivos com gráficos de criticidade."""

    # ── Cores ReportLab ───────────────────────────────────────────────
    COLOR_PRIMARY   = colors.HexColor('#1e3a8a')
    COLOR_SECONDARY = colors.HexColor('#3b82f6')
    COLOR_ACCENT    = colors.HexColor('#60a5fa')
    COLOR_CRITICAL  = colors.HexColor('#dc2626')
    COLOR_HIGH      = colors.HexColor('#ea580c')
    COLOR_MEDIUM    = colors.HexColor('#ca8a04')
    COLOR_LOW       = colors.HexColor('#16a34a')
    COLOR_HEADER    = colors.HexColor('#0f172a')
    COLOR_BG_LIGHT  = colors.HexColor('#f8fafc')

    PROVIDER_NAMES = {
        'AWS_S3':      'AWS S3',
        'AWS_IAM':     'AWS IAM (CSPM)',
        'GCP_IAM':     'GCP IAM (CSPM)',
        'AZURE_IAM':   'Azure IAM / Entra ID (CSPM)',
        'GCS':         'Google Cloud Storage',
        'AZURE':       'Azure Blob Storage',
        'AZURE_BLOB':  'Azure Blob Storage',
        'KUBERNETES':  'Kubernetes Cluster',
        'UNIVERSAL':   'Multi-Cloud',
    }

    SENSITIVE_EXTENSIONS = {
        '.env', '.pem', '.key', '.crt', '.cer', '.p12', '.pfx',
        '.sql', '.db', '.sqlite', '.mdb', '.bak', '.dump',
        '.yaml', '.yml', '.ini', '.conf', '.cfg',
        '.zip', '.tar', '.gz', '.rar', '.7z',
        '.py', '.php', '.rb', '.java', '.go',
        '.json', '.xml', '.csv',
    }

    SEV_BASE = {'CRITICAL': 87, 'HIGH': 62, 'MEDIUM': 37, 'LOW': 12}

    def __init__(self, scan_data: dict, client_info: dict = None):
        print("🔧 Inicializando gerador de relatórios...")
        print(f"📊 Dados recebidos: {list(scan_data.keys())}")

        self._validate_scan_data(scan_data)

        self.scan_data   = scan_data
        self.client_info = client_info or {}
        self.timestamp   = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        self.report_date = datetime.now()

        self.output_dir = Path("./reports_executive")
        self.output_dir.mkdir(exist_ok=True)

        # Diretório temporário para gráficos
        self.charts_dir = self.output_dir / "charts_tmp"
        self.charts_dir.mkdir(exist_ok=True)
        self.charts = ChartEngine(self.charts_dir)

        self.provider_name = self._get_provider_name()
        print(f"☁️  Provider: {self.provider_name}")

        # Análises de dados
        self.vulnerabilities    = self._safe(self._analyze_vulnerabilities,  [])
        self.recommendations    = self._safe(self._generate_recommendations,  [])
        self.risk_level         = self._safe(self._calculate_risk_level,      self._default_risk())
        self.compliance_status  = self._safe(self._assess_compliance,         [])
        self.extension_stats    = self._safe(self._analyze_extensions,        [])
        self.size_distribution  = self._safe(self._analyze_size_distribution, {})
        self.top_critical_files = self._safe(self._get_top_critical_files,    [])
        self.top_largest_files  = self._safe(self._get_top_largest_files,     [])
        self.previous_scan      = self._safe(self._load_previous_scan,        None)
        self._safe(self._save_current_scan_snapshot, None)

        print("📈 Gerando gráficos...")
        self._generate_all_charts()

    # ── Helpers de inicialização ───────────────────────────────────────
    def _safe(self, fn, default):
        try:
            return fn()
        except Exception as e:
            print(f"⚠️  {fn.__name__}: {e}")
            return default

    def _default_risk(self):
        return {'level':'DESCONHECIDO','score':0,'color':self.COLOR_MEDIUM,
                'action':'ANÁLISE NECESSÁRIA','critical_count':0,'high_count':0,'medium_count':0}

    def _validate_scan_data(self, sd):
        if not sd: raise ValueError("scan_data vazio")
        # IAM usa 'target' (account/123456) em vez de 'bucket'
        if 'bucket' not in sd:
            sd['bucket'] = sd.get('target', sd.get('account_id', 'unknown'))
        sd.setdefault('files', [])
        sd.setdefault('severity_distribution', self._calc_sev_dist(sd.get('files', [])))
        sd.setdefault('risk_score', 0)
        sd.setdefault('provider', 'UNIVERSAL')
        print("✅ Validação concluída")

    def _calc_sev_dist(self, files):
        dist = {'critical':0,'high':0,'medium':0,'low':0}
        for f in files:
            k = f.get('severity','LOW').lower()
            if k in dist: dist[k] += 1
        return dist

    def _get_provider_name(self):
        p = self.scan_data.get('provider','UNIVERSAL').upper()
        return self.PROVIDER_NAMES.get(p, p)

    # ── Análises ──────────────────────────────────────────────────────
    def _calculate_risk_level(self):
        sd    = self.scan_data.get('severity_distribution', {})
        score = self.scan_data.get('risk_score', 0)
        crit, high, med = sd.get('critical',0), sd.get('high',0), sd.get('medium',0)
        if crit > 0:    lvl, col, act = 'CRÍTICO', self.COLOR_CRITICAL, 'AÇÃO IMEDIATA REQUERIDA'
        elif high > 5:  lvl, col, act = 'ALTO',    self.COLOR_HIGH,     'AÇÃO URGENTE REQUERIDA'
        elif high > 0 or med > 10:
                        lvl, col, act = 'MÉDIO',   self.COLOR_MEDIUM,   'AÇÃO NECESSÁRIA'
        else:           lvl, col, act = 'BAIXO',   self.COLOR_LOW,      'MONITORAMENTO RECOMENDADO'
        return {'level':lvl,'score':score,'color':col,'action':act,
                'critical_count':crit,'high_count':high,'medium_count':med}

    def _assess_compliance(self):
        sd   = self.scan_data.get('severity_distribution', {})
        prov = self.scan_data.get('provider','').upper()
        crit = sd.get('critical', 0)
        high = sd.get('high', 0)
        out  = []

        if prov == 'AWS_IAM':
            findings = self.scan_data.get('files', [])
            checks   = [f.get('check','') for f in findings]
            root_mfa   = self.scan_data.get('summary', {}).get('root_mfa', True)
            no_mfa     = any('MFA' in c for c in checks)
            old_keys   = any('rotação' in c or 'Key' in c for c in checks)
            admin      = any('AdministratorAccess' in c for c in checks)
            wildcard   = any('wildcard' in c.lower() or 'Action=*' in c for c in checks)
            out.append({'name':'CIS AWS Benchmark v1.4',
                        'status':'❌ NÃO CONFORME' if crit > 0 else ('⚠️ PARCIAL' if high > 0 else '✅ CONFORME'),
                        'issues': f'{crit+high} controles CIS violados (root MFA, key rotation, admin access)'})
            out.append({'name':'ISO 27001 — A.9 (Controle de Acesso)',
                        'status':'❌ NÃO CONFORME' if (admin or wildcard) else ('⚠️ PARCIAL' if no_mfa else '✅ CONFORME'),
                        'issues':'Permissões excessivas ou MFA ausente' if (admin or no_mfa) else 'Revisar políticas de acesso'})
            out.append({'name':'NIST SP 800-53 — IA (Identificação e Auth.)',
                        'status':'❌ NÃO CONFORME' if (no_mfa or not root_mfa) else '✅ CONFORME',
                        'issues':'MFA não universal — controles IA-2 e IA-5 violados' if no_mfa else 'Autenticação em conformidade'})
            out.append({'name':'AWS Well-Architected — Security Pillar',
                        'status':'⚠️ PARCIAL' if (old_keys or no_mfa) else ('✅ CONFORME' if crit == 0 else '❌ NÃO CONFORME'),
                        'issues':'Rotação de chaves e MFA são requisitos do pilar de segurança' if old_keys else 'Revisar periodicamente'})
            return out

        if prov == 'GCP_IAM':
            findings = self.scan_data.get('files', [])
            checks   = [f.get('check','') for f in findings]
            public   = any('público' in c.lower() for c in checks)
            primitive = any('primitive' in c.lower() or 'owner' in c.lower() or 'editor' in c.lower() for c in checks)
            sa_keys  = any('SA Key' in c or 'rotação' in c for c in checks)
            out.append({'name':'CIS Google Cloud Foundations Benchmark v2.0',
                        'status':'❌ NÃO CONFORME' if crit > 0 else ('⚠️ PARCIAL' if high > 0 else '✅ CONFORME'),
                        'issues':f'{crit+high} controles CIS violados (public bindings, primitive roles, SA key rotation)'})
            out.append({'name':'ISO 27001 — A.9 (Controle de Acesso)',
                        'status':'❌ NÃO CONFORME' if (public or primitive) else ('⚠️ PARCIAL' if high > 0 else '✅ CONFORME'),
                        'issues':'Bindings públicos ou primitive roles violam princípio do menor privilégio' if (public or primitive) else 'Revisar políticas de acesso'})
            out.append({'name':'NIST SP 800-53 — AC / IA (Acesso e Identidade)',
                        'status':'❌ NÃO CONFORME' if public else ('⚠️ PARCIAL' if primitive else '✅ CONFORME'),
                        'issues':'Acesso público não autorizado detectado — controles AC-3 e IA-2 violados' if public else 'Revisar princípio do menor privilégio'})
            out.append({'name':'Google Cloud Security Foundations Guide',
                        'status':'⚠️ PARCIAL' if (sa_keys or primitive) else ('✅ CONFORME' if crit == 0 else '❌ NÃO CONFORME'),
                        'issues':'Workload Identity Federation recomendado para eliminar SA keys' if sa_keys else 'Revisar periodicamente'})
            return out

        if prov == 'AZURE_IAM':
            findings = self.scan_data.get('files', [])
            checks   = [f.get('check','') for f in findings]
            no_mfa   = any('MFA' in c for c in checks)
            guest_adm = any('Guest' in c for c in checks)
            wildcard = any('wildcard' in c.lower() for c in checks)
            sp_exp   = any('expirada' in c.lower() or 'expira' in c.lower() for c in checks)
            out.append({'name':'CIS Microsoft Azure Foundations Benchmark v2.0',
                        'status':'❌ NÃO CONFORME' if crit > 0 else ('⚠️ PARCIAL' if high > 0 else '✅ CONFORME'),
                        'issues':f'{crit+high} controles CIS violados (MFA, Global Admins, Guest roles, SP credentials)'})
            out.append({'name':'ISO 27001 — A.9 (Controle de Acesso)',
                        'status':'❌ NÃO CONFORME' if (no_mfa or guest_adm) else ('⚠️ PARCIAL' if wildcard else '✅ CONFORME'),
                        'issues':'MFA ausente ou Guest com acesso privilegiado' if (no_mfa or guest_adm) else 'Revisar custom roles'})
            out.append({'name':'NIST SP 800-53 — IA-2 / IA-5 (Autenticação)',
                        'status':'❌ NÃO CONFORME' if no_mfa else '✅ CONFORME',
                        'issues':'Conditional Access MFA não enforced para todos os usuários' if no_mfa else 'MFA em conformidade'})
            out.append({'name':'Microsoft Cloud Security Benchmark (MCSB)',
                        'status':'⚠️ PARCIAL' if (sp_exp or wildcard) else ('✅ CONFORME' if crit == 0 else '❌ NÃO CONFORME'),
                        'issues':'Service Principal credentials expiradas ou custom roles com wildcard' if (sp_exp or wildcard) else 'Revisar periodicamente'})
            return out

        if crit > 0 or high > 0:
            out.append({'name':'LGPD / GDPR','status':'❌ NÃO CONFORME','issues':'Dados pessoais potencialmente expostos'})
        else:
            out.append({'name':'LGPD / GDPR','status':'⚠️ REVISAR','issues':'Verificar classificação de dados'})
        out.append({'name':'ISO 27001',
                    'status': '❌ NÃO CONFORME' if crit > 0 else '⚠️ PARCIAL',
                    'issues': 'Controles de acesso inadequados' if crit > 0 else 'Revisar políticas'})
        if any('card' in f.get('key','').lower() or 'payment' in f.get('key','').lower()
               for f in self.scan_data.get('files',[])):
            out.append({'name':'PCI DSS','status':'❌ CRÍTICO','issues':'Dados de pagamento expostos'})
        return out

    def _analyze_vulnerabilities(self):
        provider = self.scan_data.get('provider', '').upper()

        if provider == 'AWS_IAM':
            iam_cats = {
                'mfa':       {'name':'MFA / Autenticação',        'icon':'[MFA]', 'items':[]},
                'keys':      {'name':'Access Keys',                'icon':'[KEY]', 'items':[]},
                'admin':     {'name':'Permissões Excessivas',      'icon':'[ADM]', 'items':[]},
                'trust':     {'name':'Cross-Account Trust',        'icon':'[TRU]', 'items':[]},
                'root':      {'name':'Conta Root / Política',      'icon':'[ROT]', 'items':[]},
                'inactive':  {'name':'Entidades Inativas',         'icon':'[INV]', 'items':[]},
            }
            for f in self.scan_data.get('files', []):
                sev   = f.get('severity', 'LOW')
                check = f.get('check', '')
                etype = f.get('entity_type', '')
                item  = {'file': f.get('key', f.get('name','unknown')),
                         'size': 0, 'severity': sev,
                         'reason': f.get('reason',''), 'last_modified': 'N/A'}
                if 'MFA' in check or 'senha' in check.lower() or 'password' in check.lower():
                    iam_cats['mfa']['items'].append(item)
                elif 'Key' in check or 'key' in check.lower():
                    iam_cats['keys']['items'].append(item)
                elif 'AdministratorAccess' in check or 'wildcard' in check.lower() or 'Action=*' in check:
                    iam_cats['admin']['items'].append(item)
                elif 'Trust' in check or 'Principal' in check:
                    iam_cats['trust']['items'].append(item)
                elif etype in ('root', 'account'):
                    iam_cats['root']['items'].append(item)
                elif 'inativo' in check.lower() or 'nunca' in check.lower():
                    iam_cats['inactive']['items'].append(item)
                else:
                    iam_cats['root']['items'].append(item)
            return [{'name':v['name'],'icon':v['icon'],'count':len(v['items']),'items':v['items'][:20]}
                    for v in iam_cats.values() if v['items']]

        if provider == 'GCP_IAM':
            gcp_cats = {
                'public':   {'name':'Bindings Públicos (allUsers)', 'icon':'[PUB]', 'items':[]},
                'primitive': {'name':'Primitive Roles (owner/editor)', 'icon':'[PRM]', 'items':[]},
                'sa_keys':  {'name':'SA Keys / Rotação',            'icon':'[KEY]', 'items':[]},
                'owners':   {'name':'Project Owners / Governança',  'icon':'[OWN]', 'items':[]},
                'highrisk': {'name':'Roles de Alto Risco',          'icon':'[ADM]', 'items':[]},
            }
            for f in self.scan_data.get('files', []):
                sev   = f.get('severity', 'LOW')
                check = f.get('check', '')
                item  = {'file': f.get('key', f.get('name','unknown')),
                         'size': 0, 'severity': sev,
                         'reason': f.get('reason',''), 'last_modified': 'N/A'}
                if 'público' in check.lower() or 'allUsers' in check or 'allAuthenticated' in check:
                    gcp_cats['public']['items'].append(item)
                elif 'primitive' in check.lower() or 'owner' in check.lower() or 'editor' in check.lower():
                    if 'SA Key' in check or 'key' in check.lower():
                        gcp_cats['sa_keys']['items'].append(item)
                    elif 'owners' in check.lower():
                        gcp_cats['owners']['items'].append(item)
                    else:
                        gcp_cats['primitive']['items'].append(item)
                elif 'SA Key' in check or 'key' in check.lower() or 'rotação' in check.lower() or 'dias' in check.lower():
                    gcp_cats['sa_keys']['items'].append(item)
                elif 'owners' in check.lower() or 'project owners' in check.lower():
                    gcp_cats['owners']['items'].append(item)
                else:
                    gcp_cats['highrisk']['items'].append(item)
            return [{'name':v['name'],'icon':v['icon'],'count':len(v['items']),'items':v['items'][:20]}
                    for v in gcp_cats.values() if v['items']]

        if provider == 'AZURE_IAM':
            az_cats = {
                'mfa':     {'name':'MFA / Conditional Access',   'icon':'[MFA]', 'items':[]},
                'admins':  {'name':'Global Admins / Owners',     'icon':'[ADM]', 'items':[]},
                'guest':   {'name':'Guest Users Privilegiados',  'icon':'[GST]', 'items':[]},
                'sp':      {'name':'Service Principals',         'icon':'[SP]',  'items':[]},
                'roles':   {'name':'Custom Roles / Wildcard',    'icon':'[ROL]', 'items':[]},
                'classic': {'name':'Classic Admins / Escopo',    'icon':'[CLS]', 'items':[]},
            }
            for f in self.scan_data.get('files', []):
                sev   = f.get('severity', 'LOW')
                check = f.get('check', '')
                etype = f.get('entity_type', '')
                item  = {'file': f.get('key', f.get('name','unknown')),
                         'size': 0, 'severity': sev,
                         'reason': f.get('reason',''), 'last_modified': 'N/A'}
                if 'MFA' in check or 'Conditional Access' in check:
                    az_cats['mfa']['items'].append(item)
                elif 'Guest' in check:
                    az_cats['guest']['items'].append(item)
                elif etype == 'service-principal' or 'expirada' in check.lower() or 'expira' in check.lower():
                    az_cats['sp']['items'].append(item)
                elif 'wildcard' in check.lower() or 'custom-role' in etype:
                    az_cats['roles']['items'].append(item)
                elif 'classic' in check.lower() or 'escopo raiz' in check.lower() or 'management group' in check.lower():
                    az_cats['classic']['items'].append(item)
                elif 'Owner' in check or 'Administrator' in check or 'Global' in check:
                    az_cats['admins']['items'].append(item)
                else:
                    az_cats['classic']['items'].append(item)
            return [{'name':v['name'],'icon':v['icon'],'count':len(v['items']),'items':v['items'][:20]}
                    for v in az_cats.values() if v['items']]

        categories = {
            'credentials': {'name':'Credenciais Expostas',    'icon':'[KEY]','items':[]},
            'databases':   {'name':'Bancos de Dados',          'icon':'[DB]', 'items':[]},
            'config':      {'name':'Arq. de Configuração',     'icon':'[CFG]','items':[]},
            'backups':     {'name':'Backups',                  'icon':'[BAK]','items':[]},
            'source_code': {'name':'Código Fonte',             'icon':'[SRC]','items':[]},
            'keys':        {'name':'Chaves Criptográficas',    'icon':'[SEC]','items':[]},
            'pii':         {'name':'Dados Pessoais (PII)',     'icon':'[PII]','items':[]},
        }
        for f in self.scan_data.get('files', []):
            key = f.get('key', f.get('name', '')).lower()
            sev = f.get('severity', 'LOW')
            if sev not in ('CRITICAL', 'HIGH'): continue
            item = {'file': f.get('key', f.get('name','unknown')),
                    'size': f.get('size', 0), 'severity': sev,
                    'reason': f.get('reason',''), 'last_modified': f.get('last_modified','N/A')}
            if any(x in key for x in ['password','credential','secret','token','api_key','apikey']):
                categories['credentials']['items'].append(item)
            elif any(x in key for x in ['.sql','.db','.sqlite','database','dump','.mdb']):
                categories['databases']['items'].append(item)
            elif any(x in key for x in ['.env','config','.ini','.yaml','.yml','.conf','settings']):
                categories['config']['items'].append(item)
            elif any(x in key for x in ['backup','.bak','.old','.zip','.tar','.gz','.rar']):
                categories['backups']['items'].append(item)
            elif any(x in key for x in ['.pem','.key','.crt','.cer','.p12','.pfx','private','rsa']):
                categories['keys']['items'].append(item)
            elif any(x in key for x in ['cpf','rg','passport','social','personal','customer']):
                categories['pii']['items'].append(item)
            elif any(x in key for x in ['.py','.java','.js','.php','.rb','.go','.cpp']):
                categories['source_code']['items'].append(item)
        return [{'name':v['name'],'icon':v['icon'],'count':len(v['items']),'items':v['items'][:20]}
                for v in categories.values() if v['items']]

    def _generate_recommendations(self):
        sd   = self.scan_data.get('severity_distribution', {})
        prov = self.scan_data.get('provider','UNIVERSAL').upper()
        crit = sd.get('critical', 0)
        high = sd.get('high', 0)
        recs = []

        if prov == 'AWS_IAM':
            if crit > 0:
                recs.append({'priority':'CRÍTICA','title':'Remediar Findings Críticos IAM Imediatamente',
                             'description':f'{crit} finding(s) CRÍTICO(s) identificados — risco imediato de comprometimento total da conta AWS.',
                             'actions':['Ativar MFA na conta root (IAM Console → Security credentials)',
                                        'Deletar access keys ativas da conta root',
                                        'Remover AdministratorAccess direta de usuários — migrar para grupos/roles',
                                        'Corrigir políticas com wildcard Action=* Resource=*',
                                        'Restringir trust policies com Principal=* em roles'],
                             'timeline':'0-24 horas','responsible':'CISO / Equipe de Segurança'})
            if high > 0:
                recs.append({'priority':'ALTA','title':'Corrigir Configurações IAM de Alto Risco',
                             'description':f'{high} finding(s) de nível ALTO — credenciais ou permissões excessivas.',
                             'actions':['Habilitar MFA para todos os usuários com acesso ao console',
                                        'Rotacionar access keys com mais de 90 dias (CIS 1.14)',
                                        'Remover ou desabilitar access keys nunca utilizadas',
                                        'Revisar e desabilitar usuários inativos há mais de 90 dias',
                                        'Implementar AWS IAM Access Analyzer para detecção contínua'],
                             'timeline':'24-72 horas','responsible':'Equipe de DevSecOps'})
            recs.append({'priority':'MÉDIA','title':'Fortalecer Política de Senhas e Governança IAM',
                         'description':'Políticas de senha fracas e ausência de rotação periódica aumentam risco de comprometimento.',
                         'actions':['Configurar MinimumPasswordLength=14 e MaxPasswordAge=90',
                                    'Habilitar PasswordReusePrevention=24',
                                    'Implementar Service Control Policies (SCPs) via AWS Organizations',
                                    'Realizar revisões trimestrais de permissões IAM'],
                         'timeline':'7-14 dias','responsible':'Equipe de Segurança / Compliance'})
            recs.append({'priority':'MÉDIA','title':'Implementar Monitoramento Contínuo de IAM',
                         'description':'Detecção proativa de mudanças e acessos anômalos em identidades é essencial para CSPM efetivo.',
                         'actions':['Habilitar AWS CloudTrail em todas as regiões para auditoria IAM',
                                    'Configurar AWS Config Rules para verificação de CIS Benchmark',
                                    'Ativar AWS GuardDuty para detecção de comportamentos anômalos',
                                    'Implementar revisão mensal de achados do AWS Security Hub'],
                         'timeline':'14-30 dias','responsible':'Equipe de SecOps'})
            return recs

        if prov == 'GCP_IAM':
            if crit > 0:
                recs.append({'priority':'CRÍTICA','title':'Remover Bindings Públicos e SA com Owner Imediatamente',
                             'description':f'{crit} finding(s) CRÍTICO(s) no projeto GCP — bindings allUsers/allAuthenticatedUsers ou SA com roles/owner.',
                             'actions':['Executar: gcloud projects get-iam-policy <PROJECT> e auditar cada binding',
                                        'Remover bindings allUsers/allAuthenticatedUsers de todas as roles',
                                        'Revogar roles/owner de Service Accounts — substituir por roles mínimas',
                                        'Ativar VPC Service Controls para impedir exfiltração',
                                        'Habilitar alertas de Audit Log para mudanças de IAM Policy'],
                             'timeline':'0-24 horas','responsible':'CISO / Equipe de Segurança GCP'})
            if high > 0:
                recs.append({'priority':'ALTA','title':'Corrigir Primitive Roles e Rotação de SA Keys',
                             'description':f'{high} finding(s) de alto risco — primitive roles e SA keys sem rotação.',
                             'actions':['Substituir roles/editor e roles/owner por roles específicas (IAM Recommender)',
                                        'Rotacionar SA keys com mais de 90 dias (CIS GCP 4.3)',
                                        'Eliminar user-managed keys — adotar Workload Identity Federation',
                                        'Criar SAs dedicadas por workload com permissões mínimas',
                                        'Auditar Cloud Audit Logs de criação de service account keys'],
                             'timeline':'24-72 horas','responsible':'Equipe de DevSecOps GCP'})
            recs.append({'priority':'MÉDIA','title':'Reduzir Project Owners e Implantar Governança',
                         'description':'Excesso de owners aumenta superfície de comprometimento no projeto GCP.',
                         'actions':['Limitar Project Owners a no máximo 2 identidades (CIS GCP 1.2)',
                                    'Ativar Organization Policies para restringir domínios de identidade',
                                    'Usar IAM Conditions para acesso granular baseado em contexto',
                                    'Realizar revisões trimestrais de permissões no projeto'],
                         'timeline':'7-14 dias','responsible':'Equipe de Segurança / Compliance GCP'})
            recs.append({'priority':'MÉDIA','title':'Implementar Monitoramento e Auditoria GCP IAM',
                         'description':'Visibilidade contínua sobre mudanças de IAM é essencial para CSPM em GCP.',
                         'actions':['Ativar Cloud Audit Logs (Data Access) em todos os serviços críticos',
                                    'Configurar Security Command Center para alertas de IAM',
                                    'Usar IAM Recommender semanalmente para detectar permissões excessivas',
                                    'Implementar VPC Service Controls para isolamento de projetos sensíveis'],
                         'timeline':'14-30 dias','responsible':'Equipe de SecOps GCP'})
            return recs

        if prov == 'AZURE_IAM':
            if crit > 0:
                recs.append({'priority':'CRÍTICA','title':'Forçar MFA e Remover Guest Admins Imediatamente',
                             'description':f'{crit} finding(s) CRÍTICO(s) no tenant Azure — ausência de MFA enforcement ou Guest com Global Admin role.',
                             'actions':['Criar Conditional Access Policy: Users=All, Grant=Require MFA (Entra ID)',
                                        'Remover todos os usuários Guest da role Global Administrator',
                                        'Ativar Privileged Identity Management (PIM) para acesso just-in-time',
                                        'Revogar sessões de usuários Guest com acesso privilegiado',
                                        'Habilitar Microsoft Defender for Identity para detecção de ameaças'],
                             'timeline':'0-24 horas','responsible':'CISO / Equipe de Segurança Azure'})
            if high > 0:
                recs.append({'priority':'ALTA','title':'Reduzir Global Admins e Renovar Credenciais SP',
                             'description':f'{high} finding(s) de alto risco — Global Admins excessivos ou SP credentials expiradas.',
                             'actions':['Limitar Global Admins a no máximo 4 identidades (CIS Azure)',
                                        'Migrar admins para PIM — acesso just-in-time com aprovação',
                                        'Renovar Service Principal credentials expiradas (Azure Key Vault)',
                                        'Remover Owners excessivos da subscription',
                                        'Auditar Entra ID Sign-in logs para comportamentos anômalos'],
                             'timeline':'24-72 horas','responsible':'Equipe de DevSecOps Azure'})
            recs.append({'priority':'MÉDIA','title':'Eliminar Wildcard Roles e Classic Admins',
                         'description':'Custom roles com wildcard e Co-Administrators aumentam superfície de ataque no Azure.',
                         'actions':['Substituir ações wildcard em custom roles por ações específicas do Resource Provider',
                                    'Remover todos os Classic Administrators (Co-Admin) — migrar para RBAC',
                                    'Revisar atribuições em escopo de Management Group/Root',
                                    'Implementar Azure Policy com iniciativa CIS Azure Foundations Benchmark'],
                         'timeline':'7-14 dias','responsible':'Equipe de Segurança / Compliance Azure'})
            recs.append({'priority':'MÉDIA','title':'Implementar Monitoramento Contínuo de Postura Azure',
                         'description':'Visibilidade contínua sobre identidades e acessos é essencial para CSPM no Azure.',
                         'actions':['Ativar Microsoft Defender for Cloud — plano de Identidade',
                                    'Configurar Microsoft Sentinel para alertas de IAM e Entra ID',
                                    'Habilitar Entra ID Identity Protection para detectar identidades comprometidas',
                                    'Realizar revisões mensais de Conditional Access e role assignments'],
                         'timeline':'14-30 dias','responsible':'Equipe de SecOps Azure'})
            return recs

        if 'KUBERNETES' in prov:
            if crit > 0:
                recs.append({'priority':'CRÍTICA','title':'Remover Credenciais Expostas em Secrets',
                             'description':f'{crit} Secret(s) com credenciais de storage em texto plano detectados.',
                             'actions':["Rotacionar imediatamente todas as credenciais expostas",
                                        "Migrar segredos para um Vault (HashiCorp Vault ou AWS Secrets Manager)",
                                        "Revogar tokens e chaves comprometidas no provedor cloud",
                                        "Auditar quem teve acesso aos Secrets afetados",
                                        "Ativar criptografia de Secrets em repouso no etcd"],
                             'timeline':'0-24 horas','responsible':'Equipe de Segurança / SRE'})
            if high > 0:
                recs.append({'priority':'ALTA','title':'Corrigir PersistentVolumes Inseguros',
                             'description':f'{high} PV(s)/PVC(s) com configuração de risco (hostPath, ReadWriteMany ou Retain).',
                             'actions':["Remover ou substituir volumes hostPath por CSI drivers seguros",
                                        "Alterar ReclaimPolicy de Retain para Delete onde aplicável",
                                        "Restringir AccessMode ReadWriteMany a workloads que realmente precisam",
                                        "Auditar permissões RBAC sobre PersistentVolumes",
                                        "Implementar PodSecurityAdmission para bloquear hostPath"],
                             'timeline':'24-72 horas','responsible':'Equipe de Infraestrutura / SRE'})
            recs.append({'priority':'MÉDIA','title':'Mover Dados Sensíveis de ConfigMaps para Secrets',
                         'description':'ConfigMaps não são criptografados; dados sensíveis devem estar em Secrets ou Vault.',
                         'actions':["Identificar todos os ConfigMaps com chaves sensíveis",
                                    "Migrar para Secrets Kubernetes ou solução de Vault externa",
                                    "Aplicar RBAC restritivo nos Secrets criados",
                                    "Habilitar auditoria no Kubernetes API Server para Secrets",
                                    "Rever pipelines CI/CD que injetam segredos via ConfigMap"],
                         'timeline':'7-14 dias','responsible':'Equipe de DevSecOps'})
            recs.append({'priority':'MÉDIA','title':'Fortalecer RBAC e Auditoria do Cluster',
                         'description':'Controle de acesso granular e logs de auditoria são essenciais para conformidade.',
                         'actions':["Aplicar princípio do menor privilégio em todos os ServiceAccounts",
                                    "Habilitar Kubernetes Audit Logs para operações em Secrets e PVs",
                                    "Revisar ClusterRoleBindings com permissões excessivas",
                                    "Implementar Network Policies para isolar namespaces",
                                    "Realizar revisões trimestrais de permissões RBAC"],
                         'timeline':'14-30 dias','responsible':'CISO / Equipe de Plataforma'})
        else:
            if crit > 0 or high > 0:
                actions_map = {
                    'AWS_S3': ["Bloquear acesso público via 'Block Public Access'",
                               "Revisar bucket policies e ACLs no console AWS",
                               "Implementar bucket encryption (AES-256 ou KMS)",
                               "Habilitar versioning para recovery",
                               "Configurar logging para auditoria"],
                    'GCS':    ["Remover 'allUsers' e 'allAuthenticatedUsers' das permissões",
                               "Ativar 'Uniform bucket-level access'",
                               "Implementar IAM Conditions",
                               "Habilitar 'Object versioning'",
                               "Configurar 'Audit logs'"],
                }
                if 'AZURE' in prov:
                    actions = ["Revisar IAM no Azure Portal",
                               "Configurar 'Private endpoints'",
                               "Habilitar 'SAS' com expiração",
                               "Implementar Azure AD",
                               "Configurar 'Network rules'"]
                else:
                    actions = actions_map.get(prov, ["Revisar políticas de acesso",
                                                      "Remover permissões públicas",
                                                      "Implementar autenticação forte",
                                                      "Habilitar criptografia em repouso",
                                                      "Configurar logs de auditoria"])
                recs.append({'priority':'CRÍTICA','title':'Restringir Acesso Público ao Storage',
                             'description':f'{crit+high} arquivo(s) com exposição pública detectados.',
                             'actions':actions,'timeline':'0-24 horas',
                             'responsible':'Equipe de Segurança / DevOps'})

            recs.append({'priority':'ALTA','title':'Implementar Criptografia de Dados',
                         'description':'Garantir criptografia em repouso e em trânsito.',
                         'actions':["Habilitar criptografia server-side",
                                    "Implementar HTTPS/TLS para transferências",
                                    "Rotacionar chaves regularmente",
                                    "Usar CMK quando possível",
                                    "Documentar gestão de chaves"],
                         'timeline':'7-14 dias','responsible':'Equipe de Segurança'})

            recs.append({'priority':'MÉDIA','title':'Implementar Monitoramento Contínuo',
                         'description':'Detectar acessos suspeitos e mudanças de configuração.',
                         'actions':["Configurar alertas para acessos anômalos",
                                    "Implementar SIEM para análise de logs",
                                    "Criar dashboards de segurança",
                                    "Estabelecer processo de resposta a incidentes",
                                    "Realizar auditorias regulares"],
                         'timeline':'14-30 dias','responsible':'Equipe de SecOps'})

            recs.append({'priority':'MÉDIA','title':'Estabelecer Políticas de Governança',
                         'description':'Classificação e proteção de dados com políticas claras.',
                         'actions':["Classificar dados por sensibilidade",
                                    "Definir políticas de retenção",
                                    "Implementar DLP",
                                    "Treinar equipes",
                                    "Revisões trimestrais de acessos"],
                         'timeline':'30-60 dias','responsible':'CISO / Compliance'})
        return recs

    def _analyze_extensions(self):
        ext_map = {}
        for f in self.scan_data.get('files', []):
            ext = Path(f.get('key', f.get('name',''))).suffix.lower() or '(sem extensão)'
            sev = f.get('severity','LOW').lower()
            if ext not in ext_map:
                ext_map[ext] = {'total':0,'critical':0,'high':0,'medium':0,'low':0,
                                'sensitive': ext in self.SENSITIVE_EXTENSIONS}
            ext_map[ext]['total'] += 1
            if sev in ext_map[ext]: ext_map[ext][sev] += 1
        result = [{'ext':k,**v} for k,v in ext_map.items()]
        result.sort(key=lambda x: (x['critical']+x['high'], x['total']), reverse=True)
        return result[:25]

    def _analyze_size_distribution(self):
        stats = {s:{'count':0,'total_bytes':0} for s in ['critical','high','medium','low']}
        for f in self.scan_data.get('files', []):
            sev  = f.get('severity','LOW').lower()
            size = int(f.get('size', 0))
            if sev in stats:
                stats[sev]['count']       += 1
                stats[sev]['total_bytes'] += size
        for k in stats:
            c = stats[k]['count']
            stats[k]['avg_bytes'] = stats[k]['total_bytes'] // c if c else 0
        stats['_total_bytes'] = sum(v['total_bytes'] for v in stats.values())
        return stats

    def _get_top_critical_files(self, n=15):
        order = {'CRITICAL':0,'HIGH':1,'MEDIUM':2,'LOW':3}
        files = sorted(self.scan_data.get('files',[]),
                       key=lambda f: (order.get(f.get('severity','LOW'),4),
                                      -int(f.get('size',0))))
        return [{'file': f.get('key',f.get('name','unknown')),
                 'severity': f.get('severity','LOW'),
                 'size': int(f.get('size',0)),
                 'reason': f.get('reason','Exposição pública'),
                 'last_modified': f.get('last_modified','N/A')}
                for f in files[:n]
                if f.get('severity','LOW') in ('CRITICAL','HIGH','MEDIUM')]

    def _get_top_largest_files(self, n=10):
        files = sorted(self.scan_data.get('files',[]),
                       key=lambda f: int(f.get('size',0)), reverse=True)
        return [{'file': f.get('key',f.get('name','unknown')),
                 'severity': f.get('severity','LOW'),
                 'size': int(f.get('size',0))} for f in files[:n]]

    def _load_previous_scan(self):
        snap = self.output_dir / f"snapshot_{self.scan_data.get('bucket','').replace('.','_')}.json"
        if snap.exists():
            with open(snap) as fh: return json.load(fh)
        return None

    def _save_current_scan_snapshot(self):
        snap = {'timestamp': self.timestamp,
                'severity_distribution': self.scan_data.get('severity_distribution',{}),
                'total_files': len(self.scan_data.get('files',[])),
                'risk_score':  self.scan_data.get('risk_score',0)}
        path = self.output_dir / f"snapshot_{self.scan_data.get('bucket','').replace('.','_')}.json"
        with open(path,'w') as fh: json.dump(snap, fh, indent=2)

    def _build_comparison(self):
        if not self.previous_scan: return None
        curr = self.scan_data.get('severity_distribution',{})
        prev = self.previous_scan.get('severity_distribution',{})
        return {'prev_timestamp': self.previous_scan.get('timestamp','N/A'),
                'prev_total':     self.previous_scan.get('total_files',0),
                'curr_total':     len(self.scan_data.get('files',[])),
                'delta_critical': curr.get('critical',0) - prev.get('critical',0),
                'delta_high':     curr.get('high',0)     - prev.get('high',0),
                'delta_medium':   curr.get('medium',0)   - prev.get('medium',0),
                'delta_score':    self.scan_data.get('risk_score',0) - self.previous_scan.get('risk_score',0)}

    # ── Narrativa executiva auto-gerada ───────────────────────────────
    def _generate_narrative(self) -> str:
        provider_raw = self.scan_data.get('provider','').upper()
        sd      = self.scan_data.get('severity_distribution', {})
        total   = len(self.scan_data.get('files', []))
        crit    = sd.get('critical', 0)
        high    = sd.get('high', 0)
        med     = sd.get('medium', 0)
        low     = sd.get('low', 0)
        prov    = self.provider_name
        bucket  = self.scan_data.get('bucket', 'o ambiente auditado')
        score   = self.scan_data.get('risk_score', 0)
        risk    = self.risk_level.get('level', 'DESCONHECIDO')
        total_b = self.size_distribution.get('_total_bytes', 0)
        date_s  = self.report_date.strftime('%d/%m/%Y às %H:%M')

        if provider_raw == 'AWS_IAM':
            summ = self.scan_data.get('summary', {})
            users_total  = summ.get('users_total', '?')
            users_no_mfa = summ.get('users_no_mfa', 0)
            keys_old     = summ.get('keys_old', 0)
            root_mfa     = summ.get('root_mfa', True)
            opening = (
                f"Em {date_s}, foi realizado um scan CSPM de IAM na conta AWS <b>{bucket}</b>, "
                f"analisando <b>{users_total} usuário(s)</b> e identificando <b>{total} finding(s)</b> de segurança. "
            )
            if crit > 0:
                risk_sentence = (
                    f"Foram identificados <b>{crit} finding(s) CRÍTICO(s)</b> e <b>{high} de nível ALTO</b>, "
                    f"indicando postura IAM de risco <b>{risk}</b> — ação imediata requerida. "
                )
            elif high > 0:
                risk_sentence = (
                    f"Não foram encontrados findings críticos, porém <b>{high} finding(s)</b> "
                    f"apresentam risco ALTO (MFA ausente, keys não rotacionadas ou permissões excessivas). "
                )
            else:
                risk_sentence = (
                    f"Não foram identificados findings críticos ou de alto risco. "
                    f"A postura IAM apresenta nível de risco <b>{risk}</b>. "
                )
            details = []
            if not root_mfa: details.append("conta root sem MFA")
            if users_no_mfa: details.append(f"{users_no_mfa} usuário(s) sem MFA")
            if keys_old:     details.append(f"{keys_old} access key(s) sem rotação")
            detail_sentence = f"Destaques: <b>{', '.join(details)}</b>. " if details else ""
            score_sentence  = (
                f"O Risk Score IAM global é de <b>{score}/100</b>, "
                f"calculado com base na severidade e volume de findings identificados. "
            )
            action = ("Ação imediata recomendada para remediar os findings críticos." if crit > 0
                      else "Priorizar as ações corretivas nos próximos 7 dias." if high > 0
                      else "Manter monitoramento contínuo e revisões periódicas de IAM.")
            return opening + risk_sentence + detail_sentence + score_sentence + action

        if provider_raw == 'GCP_IAM':
            summ        = self.scan_data.get('summary', {})
            project_id  = summ.get('project_id', bucket)
            sa_keys_old = summ.get('sa_keys_old', 0)
            pub_bind    = summ.get('public_bindings', 0)
            opening = (
                f"Em {date_s}, foi realizado um scan CSPM de IAM no projeto GCP <b>{project_id}</b>, "
                f"identificando <b>{total} finding(s)</b> de segurança em políticas IAM e Service Accounts. "
            )
            if crit > 0:
                risk_sentence = (
                    f"Foram identificados <b>{crit} finding(s) CRÍTICO(s)</b> — incluindo bindings públicos "
                    f"(allUsers/allAuthenticatedUsers) ou Service Accounts com roles/owner, "
                    f"indicando postura de risco <b>{risk}</b> com exposição imediata. "
                )
            elif high > 0:
                risk_sentence = (
                    f"Não foram encontrados findings críticos, porém <b>{high} finding(s)</b> "
                    f"apresentam risco ALTO (primitive roles, SA keys sem rotação ou roles de alto risco). "
                )
            else:
                risk_sentence = f"Postura IAM GCP com nível de risco <b>{risk}</b>. "
            details = []
            if pub_bind:    details.append(f"{pub_bind} binding(s) público(s)")
            if sa_keys_old: details.append(f"{sa_keys_old} SA key(s) sem rotação")
            detail_sentence = f"Destaques: <b>{', '.join(details)}</b>. " if details else ""
            score_sentence  = (
                f"Risk Score IAM GCP: <b>{score}/100</b>, "
                f"calculado pela severidade dos findings identificados no projeto. "
            )
            action = ("Ação imediata recomendada — remover bindings públicos e SA com roles/owner." if crit > 0
                      else "Priorizar rotação de SA keys e remoção de primitive roles nos próximos 7 dias." if high > 0
                      else "Manter monitoramento via Security Command Center e IAM Recommender.")
            return opening + risk_sentence + detail_sentence + score_sentence + action

        if provider_raw == 'AZURE_IAM':
            summ      = self.scan_data.get('summary', {})
            sub_id    = summ.get('subscription_id', bucket)
            no_mfa    = summ.get('no_mfa_enforcement', False)
            guest_adm = summ.get('guest_admins', 0)
            opening = (
                f"Em {date_s}, foi realizado um scan CSPM de IAM na subscription Azure <b>{sub_id}</b> "
                f"e tenant Entra ID, identificando <b>{total} finding(s)</b> de segurança. "
            )
            if crit > 0:
                risk_sentence = (
                    f"Foram identificados <b>{crit} finding(s) CRÍTICO(s)</b> — "
                    f"{'ausência de MFA enforcement via Conditional Access' if no_mfa else ''}"
                    f"{' e ' if no_mfa and guest_adm else ''}"
                    f"{f'{guest_adm} usuário(s) Guest com roles privilegiadas' if guest_adm else ''}, "
                    f"indicando postura de risco <b>{risk}</b>. "
                )
            elif high > 0:
                risk_sentence = (
                    f"Não foram encontrados findings críticos, porém <b>{high} finding(s)</b> "
                    f"apresentam risco ALTO (Global Admins excessivos, SP credentials expiradas ou wildcard roles). "
                )
            else:
                risk_sentence = f"Postura IAM Azure / Entra ID com nível de risco <b>{risk}</b>. "
            details = []
            if no_mfa:    details.append("MFA não enforced para todos os usuários")
            if guest_adm: details.append(f"{guest_adm} Guest(s) com acesso privilegiado")
            detail_sentence = f"Destaques: <b>{', '.join(details)}</b>. " if details else ""
            score_sentence  = (
                f"Risk Score IAM Azure: <b>{score}/100</b>, "
                f"calculado pela severidade dos findings na subscription e Entra ID. "
            )
            action = ("Ação imediata — enforçar MFA via Conditional Access e remover Guest admins." if crit > 0
                      else "Priorizar redução de Global Admins e renovação de credenciais SP nos próximos 7 dias." if high > 0
                      else "Manter monitoramento via Microsoft Defender for Cloud e revisões periódicas.")
            return opening + risk_sentence + detail_sentence + score_sentence + action

        # Frase de abertura
        opening = (
            f"Em {date_s}, foi realizada uma auditoria de segurança em <b>{bucket}</b> "
            f"({prov}), resultando na análise de <b>{total:,} objeto(s)</b> com volume "
            f"total de <b>{self._format_size(total_b)}</b> de dados. "
        )

        # Avaliação de risco
        if crit > 0:
            risk_sentence = (
                f"Foram identificadas <b>{crit} exposição(ões) CRÍTICA(s)</b> e <b>{high} de nível ALTO</b>, "
                f"indicando risco <b>{risk}</b> imediato ao ambiente. "
            )
        elif high > 0:
            risk_sentence = (
                f"Não foram encontradas exposições críticas, porém <b>{high} arquivo(s)</b> "
                f"apresentam risco ALTO de exposição pública. "
            )
        else:
            risk_sentence = (
                f"Não foram identificadas exposições críticas ou de alto risco. "
                f"O ambiente apresenta nível de risco <b>{risk}</b>. "
            )

        # Score
        score_sentence = (
            f"O Risk Score global do ambiente é de <b>{score}/100</b>, "
            f"calculado com base na distribuição de severidade e volume de dados expostos. "
        )

        # Recomendação
        if crit > 0 or high > 5:
            action = "Ação imediata é recomendada para mitigar os riscos identificados dentro de 24 horas."
        elif high > 0 or med > 5:
            action = "Recomenda-se priorizar as ações corretivas nos próximos 7 dias."
        else:
            action = "Recomenda-se monitoramento contínuo e revisão das políticas de acesso."

        return opening + risk_sentence + score_sentence + action

    # ── Gerar todos os gráficos ────────────────────────────────────────
    def _generate_all_charts(self):
        files    = self.scan_data.get('files', [])
        sd       = self.scan_data.get('severity_distribution', {})
        total    = len(files)
        provider = self.scan_data.get('provider','').upper()
        vuln     = [f for f in files if f.get('severity','LOW') != 'LOW']
        gauge_files = vuln if vuln else files[:20]

        self.ch = {}
        self.ch['env_gauge'] = self._safe(lambda: self.charts.environment_gauge(
                                    self.scan_data.get('risk_score', 0),
                                    self.risk_level.get('level', 'MÉDIO')), None)
        self.ch['donut']     = self._safe(lambda: self.charts.severity_donut(sd, total), None)
        self.ch['funnel']    = self._safe(lambda: self.charts.priority_funnel(sd, total), None)
        self.ch['risk_all']  = self._safe(lambda: self.charts.risk_per_file(files[:20], 'Top 20'), None)
        self.ch['risk_vuln'] = self._safe(lambda: self.charts.risk_per_file(gauge_files, 'Vulneráveis', 'risk_vuln'), None) if vuln else self.ch['risk_all']

        if 'IAM' in provider:
            # Todos os providers IAM (AWS_IAM, GCP_IAM, AZURE_IAM) usam iam_check_bars
            self.ch['cat_bars']  = self._safe(lambda: self.charts.iam_check_bars(files), None)
            self.ch['size']      = None   # sem volume de dados em IAM
            self.ch['heatmap']   = None   # heatmap de arquivo não se aplica
            self.ch['type_pie']  = None   # sem extensões de arquivo
        else:
            self.ch['cat_bars']  = self._safe(lambda: self.charts.category_bars(files), None)
            self.ch['size']      = self._safe(lambda: self.charts.size_per_file(files[:20]), None)
            self.ch['heatmap']   = self._safe(lambda: self.charts.risk_heatmap(files[:15]), None)
            self.ch['type_pie']  = self._safe(lambda: self.charts.type_pie(files), None)

        gauge_limit = gauge_files[:20]
        if len(gauge_limit) > 9:
            max_sz_gl = max((float(f.get('size', 0)) for f in gauge_limit), default=1) or 1
            g_scores  = [self.charts._file_score(f, max_sz_gl) for f in gauge_limit]
            if g_scores and (max(g_scores) - min(g_scores)) < 6:
                gauge_limit = gauge_files[:9]
        self.ch['gauges'] = self._safe(lambda gl=gauge_limit: self.charts.all_gauges(gl), [])
        print(f"✅ {sum(1 for v in self.ch.values() if v)} gráficos gerados "
              f"({len(self.ch['gauges'])} gauges individuais)")

    # ── Helpers de formatação ──────────────────────────────────────────
    def _format_size(self, b):
        try: b = int(b)
        except: return "0 B"
        for u in ['B','KB','MB','GB']:
            if b < 1024: return f"{b:.1f} {u}"
            b /= 1024
        return f"{b:.1f} TB"

    def _calc_pct(self, v, t):
        return f"{v/t*100:.1f}" if t else "0.0"

    def _delta_str(self, v):
        return f"+{v}" if v > 0 else str(v)

    def _compute_hash(self, filepath):
        h = hashlib.sha256()
        with open(filepath,'rb') as fh:
            for chunk in iter(lambda: fh.read(8192), b''):
                h.update(chunk)
        return h.hexdigest()

    # ══════════════════════════════════════════════════════════════════
    # PDF GENERATION
    # ══════════════════════════════════════════════════════════════════
    def _create_header_footer(self, canvas, doc):
        canvas.saveState()
        # Header — barra primária + acento secundário
        canvas.setFillColor(self.COLOR_PRIMARY)
        canvas.rect(0, A4[1]-0.65*inch, A4[0], 0.65*inch, fill=True, stroke=False)
        canvas.setFillColor(self.COLOR_SECONDARY)
        canvas.rect(0, A4[1]-0.68*inch, A4[0], 0.04*inch, fill=True, stroke=False)
        canvas.setFillColor(colors.white)
        canvas.setFont('Helvetica-Bold', 11)
        canvas.drawString(0.55*inch, A4[1]-0.38*inch, "RELATÓRIO DE SEGURANÇA EXECUTIVO")
        canvas.setFont('Helvetica', 9)
        canvas.drawRightString(A4[0]-0.55*inch, A4[1]-0.38*inch,
                               self.report_date.strftime("%d/%m/%Y"))
        # Footer
        canvas.setFillColor(colors.HexColor('#f1f5f9'))
        canvas.rect(0, 0, A4[0], 0.65*inch, fill=True, stroke=False)
        canvas.setFillColor(self.COLOR_PRIMARY)
        canvas.rect(0, 0.62*inch, A4[0], 0.03*inch, fill=True, stroke=False)
        canvas.setFillColor(colors.HexColor('#64748b'))
        canvas.setFont('Helvetica', 8)
        canvas.drawString(0.55*inch, 0.25*inch,
                          f"Security Multicloud Scanner  |  {self.provider_name}")
        canvas.drawRightString(A4[0]-0.55*inch, 0.25*inch, f"Página {doc.page}")
        canvas.restoreState()

    def _create_cover_page(self, canvas, doc):
        """Capa: painel escuro superior + área de dados branca."""
        canvas.saveState()
        w, h = A4

        # Painel escuro superior (65% da altura)
        canvas.setFillColor(self.COLOR_PRIMARY)
        canvas.rect(0, h * 0.35, w, h * 0.65, fill=True, stroke=False)

        # Faixa de acento secondary
        canvas.setFillColor(self.COLOR_SECONDARY)
        canvas.rect(0, h * 0.35 - 0.06*inch, w, 0.06*inch, fill=True, stroke=False)

        # Círculos decorativos no canto superior direito
        for r, alpha_col in [(0.55*inch,'#1d4ed8'), (0.30*inch,'#2563eb'), (0.14*inch,'#3b82f6')]:
            canvas.setFillColor(colors.HexColor(alpha_col))
            canvas.circle(w - 0.5*inch, h - 0.5*inch, r, fill=True, stroke=False)

        # Título em branco
        canvas.setFillColor(colors.white)
        canvas.setFont('Helvetica-Bold', 34)
        canvas.drawCentredString(w / 2, h * 0.71, 'RELATÓRIO DE SEGURANÇA')
        canvas.setFillColor(colors.HexColor('#93c5fd'))
        canvas.setFont('Helvetica-Bold', 34)
        canvas.drawCentredString(w / 2, h * 0.64, 'EXECUTIVO')

        # Subtítulo
        canvas.setFillColor(colors.HexColor('#bfdbfe'))
        canvas.setFont('Helvetica', 12)
        canvas.drawCentredString(w / 2, h * 0.57,
                                 f'Auditoria de Storage Multicloud  ·  {self.provider_name}')

        # Linha separadora
        canvas.setStrokeColor(colors.HexColor('#3b82f6'))
        canvas.setLineWidth(1.5)
        canvas.line(1.8*inch, h * 0.52, w - 1.8*inch, h * 0.52)

        # Data centralizada
        canvas.setFillColor(colors.HexColor('#94a3b8'))
        canvas.setFont('Helvetica', 10)
        canvas.drawCentredString(w / 2, h * 0.46,
                                 self.report_date.strftime('%d/%m/%Y  ·  %H:%M'))

        # Caixa de dados do cliente
        box_x, box_w = 0.9*inch, w - 1.8*inch
        box_y, box_h = h * 0.12, h * 0.22
        canvas.setFillColor(colors.HexColor('#f8fafc'))
        canvas.rect(box_x, box_y, box_w, box_h, fill=True, stroke=False)
        canvas.setStrokeColor(colors.HexColor('#e2e8f0'))
        canvas.setLineWidth(0.8)
        canvas.rect(box_x, box_y, box_w, box_h, fill=False, stroke=True)
        # Borda esquerda colorida
        canvas.setFillColor(self.COLOR_SECONDARY)
        canvas.rect(box_x, box_y, 0.06*inch, box_h, fill=True, stroke=False)

        info_x   = box_x + 0.25*inch
        val_x    = box_x + 1.5*inch
        info_y   = box_y + box_h - 0.28*inch
        line_gap = 0.045*inch * 1.2 * 25.4 / 25.4  # espaçamento entre linhas

        rows = [
            ('Cliente:', self.client_info.get('name', '')),
            ('Contato:', self.client_info.get('contact', '')),
            ('Alvo:',    self.scan_data.get('bucket', '')),
            ('Data:',    self.report_date.strftime('%d/%m/%Y às %H:%M')),
        ]
        for label, value in rows:
            canvas.setFillColor(self.COLOR_PRIMARY)
            canvas.setFont('Helvetica-Bold', 10.5)
            canvas.drawString(info_x, info_y, label)
            canvas.setFillColor(colors.HexColor('#1e293b'))
            canvas.setFont('Helvetica', 10.5)
            canvas.drawString(val_x, info_y, str(value)[:55])
            info_y -= 0.052 * h

        # Rodapé da capa
        canvas.setFillColor(colors.HexColor('#94a3b8'))
        canvas.setFont('Helvetica', 8)
        canvas.drawCentredString(w / 2, 0.28*inch,
                                 'Security Multicloud Scanner  ·  Documento Confidencial')
        canvas.restoreState()

    def generate_pdf(self) -> str:
        print("📄 Gerando PDF...")
        bucket   = self.scan_data.get('bucket','scan').replace('.','_').replace('/','_')
        filename = f"relatorio_{bucket}_{self.timestamp}.pdf"
        filepath = self.output_dir / filename

        doc   = SimpleDocTemplate(str(filepath), pagesize=A4,
                                  topMargin=0.75*inch, bottomMargin=0.75*inch,
                                  leftMargin=0.65*inch, rightMargin=0.65*inch)
        story = []
        styles = getSampleStyleSheet()

        H1 = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=15,
                             textColor=colors.white, spaceAfter=14,
                             spaceBefore=14, fontName='Helvetica-Bold',
                             backColor=self.COLOR_PRIMARY,
                             leftPadding=10, rightPadding=10,
                             topPadding=6, bottomPadding=6, leading=20)
        H2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=11,
                             textColor=self.COLOR_PRIMARY, spaceAfter=8,
                             spaceBefore=10, fontName='Helvetica-Bold',
                             borderPad=4, leftPadding=0,
                             borderColor=self.COLOR_SECONDARY,
                             borderWidth=0)
        NOTE = ParagraphStyle('NOTE', fontSize=8, textColor=colors.HexColor('#6b7280'),
                              alignment=TA_CENTER, fontName='Helvetica-Oblique',
                              spaceAfter=4)

        def add_chart(path, width=5.8*inch, height=None, caption=None):
            if path and Path(path).exists():
                story.append(Spacer(1, 0.1*inch))
                h = height if height is not None else width * 0.55
                story.append(RLImage(path, width=width, height=h))
                if caption:
                    story.append(Paragraph(caption, NOTE))
                story.append(Spacer(1, 0.15*inch))

        sd     = self.scan_data.get('severity_distribution', {})
        total  = len(self.scan_data.get('files', []))
        risk   = self.risk_level

        # ── CAPA — desenhada inteiramente via canvas (_create_cover_page) ──
        story.append(PageBreak())

        # ── 1. RESUMO EXECUTIVO ────────────────────────────────────────
        story.append(Paragraph("1. RESUMO EXECUTIVO", H1))
        story.append(Spacer(1, 0.12*inch))

        # ── Risk Overview — título + subtítulo (estilo dashboard) ─────
        story.append(Paragraph("RISK OVERVIEW",
                                ParagraphStyle('RO_TTL', fontSize=11, fontName='Helvetica-Bold',
                                               textColor=self.COLOR_PRIMARY, spaceAfter=2)))
        story.append(Paragraph("Distribuição de exposição por severidade",
                                ParagraphStyle('RO_SUB', fontSize=8.5,
                                               textColor=colors.HexColor('#64748b'), spaceAfter=8)))

        # ── Cards de severidade — estilo Risk Overview (borda colorida superior) ──
        sev_colors = [self.COLOR_CRITICAL, self.COLOR_HIGH, self.COLOR_MEDIUM, self.COLOR_LOW]
        sev_labels = ['CRÍTICO', 'ALTO', 'MÉDIO', 'BAIXO']
        sev_keys   = ['critical', 'high', 'medium', 'low']
        sev_vals   = [sd.get(k, 0) for k in sev_keys]

        ST = ParagraphStyle
        row_lbl = [Paragraph(f"<b>{lbl}</b>",
                              ST(f'cl_{i}', fontSize=9, textColor=colors.HexColor('#64748b'),
                                 fontName='Helvetica-Bold', alignment=TA_CENTER))
                   for i, lbl in enumerate(sev_labels)]
        row_val = [Paragraph(f"<b>{val:,}</b>",
                              ST(f'cn_{i}', fontSize=30, textColor=col,
                                 fontName='Helvetica-Bold', alignment=TA_CENTER, leading=34))
                   for i, (val, col) in enumerate(zip(sev_vals, sev_colors))]
        row_pct = [Paragraph(
                       f"{val/total*100:.1f}%" if total else "0.0%",
                       ST(f'cp_{i}', fontSize=10,
                          textColor=col if val > 0 else colors.HexColor('#94a3b8'),
                          fontName='Helvetica-Bold', alignment=TA_CENTER))
                   for i, (val, col) in enumerate(zip(sev_vals, sev_colors))]

        card_row = Table([row_lbl, row_val, row_pct], colWidths=[1.58*inch]*4)
        card_row.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.white),
            # Bordas laterais e inferior finas
            ('LINEAFTER',  (0, 0), (2, -1),  0.5, colors.HexColor('#e2e8f0')),
            ('LINEBEFORE', (0, 0), (0, -1),  0.5, colors.HexColor('#e2e8f0')),
            ('LINEAFTER',  (3, 0), (3, -1),  0.5, colors.HexColor('#e2e8f0')),
            ('LINEBELOW',  (0, -1), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            # Borda superior ESPESSA colorida por coluna (Risk Overview style)
            ('LINEABOVE', (0, 0), (0, 0), 4, self.COLOR_CRITICAL),
            ('LINEABOVE', (1, 0), (1, 0), 4, self.COLOR_HIGH),
            ('LINEABOVE', (2, 0), (2, 0), 4, self.COLOR_MEDIUM),
            ('LINEABOVE', (3, 0), (3, 0), 4, self.COLOR_LOW),
            ('PADDING',       (0, 0), (-1, -1), 10),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING',    (0, 0), (-1,  0), 14),
            ('BOTTOMPADDING', (0,-1), (-1, -1), 14),
        ]))
        story.append(card_row)
        story.append(Spacer(1, 0.18*inch))

        # ── KPIs IAM adicionais ───────────────────────────────────────
        iam_prov_kpi = self.scan_data.get('provider', '').upper()
        if is_iam and iam_summ:
            def _iam_kpi(label, val, ok, unit=''):
                col  = self.COLOR_LOW if ok else self.COLOR_CRITICAL
                sval = f"{val}{unit}"
                return [Paragraph(f"<b>{label}</b>",
                                  ParagraphStyle('ik_l', fontSize=8, textColor=colors.HexColor('#64748b'),
                                                 fontName='Helvetica-Bold', alignment=TA_CENTER)),
                        Paragraph(f"<b>{sval}</b>",
                                  ParagraphStyle('ik_v', fontSize=20, textColor=col,
                                                 fontName='Helvetica-Bold', alignment=TA_CENTER, leading=24)),
                        Paragraph('✅ OK' if ok else '❌ Risco',
                                  ParagraphStyle('ik_s', fontSize=8, textColor=col,
                                                 fontName='Helvetica-Bold', alignment=TA_CENTER))]

            if iam_prov_kpi == 'AWS_IAM':
                root_mfa_ok  = iam_summ.get('root_mfa', True)
                root_keys    = iam_summ.get('root_keys_active', False)
                users_no_mfa = iam_summ.get('users_no_mfa', 0)
                keys_old     = iam_summ.get('keys_old', 0)
                admin_direct = iam_summ.get('admin_direct', 0)
                iam_kpi_rows = [
                    _iam_kpi('Root MFA',     'SIM' if root_mfa_ok else 'NÃO', root_mfa_ok),
                    _iam_kpi('Root Keys',    'NÃO' if not root_keys else 'SIM', not root_keys),
                    _iam_kpi('Sem MFA',      users_no_mfa, users_no_mfa == 0, ' users'),
                    _iam_kpi('Keys Velhas',  keys_old,     keys_old == 0,    ' keys'),
                    _iam_kpi('Admin Direto', admin_direct, admin_direct == 0, ' users'),
                ]
            elif iam_prov_kpi == 'GCP_IAM':
                pub_bind    = iam_summ.get('public_bindings', 0)
                sa_keys_old = iam_summ.get('sa_keys_old', 0)
                prm_count   = len([f for f in self.scan_data.get('files', []) if 'primitive' in f.get('check','').lower() or 'owner' in f.get('check','').lower()])
                own_count   = len([f for f in self.scan_data.get('files', []) if 'owners' in f.get('check','').lower()])
                hr_count    = len([f for f in self.scan_data.get('files', []) if 'alto risco' in f.get('check','').lower()])
                iam_kpi_rows = [
                    _iam_kpi('Bind. Públicos', pub_bind,    pub_bind == 0,    ''),
                    _iam_kpi('Prim. Roles',    prm_count,   prm_count == 0,   ''),
                    _iam_kpi('SA Keys Old',    sa_keys_old, sa_keys_old == 0, ''),
                    _iam_kpi('Owners+',        own_count,   own_count == 0,   ''),
                    _iam_kpi('High Risk',      hr_count,    hr_count == 0,    ''),
                ]
            elif iam_prov_kpi == 'AZURE_IAM':
                no_mfa    = iam_summ.get('no_mfa_enforcement', False)
                guest_adm = iam_summ.get('guest_admins', 0)
                sp_exp    = len([f for f in self.scan_data.get('files', []) if 'expirada' in f.get('check','').lower()])
                owners    = len([f for f in self.scan_data.get('files', []) if 'Owner' in f.get('check','') and 'Guest' not in f.get('check','')])
                wildcard  = len([f for f in self.scan_data.get('files', []) if 'wildcard' in f.get('check','').lower()])
                iam_kpi_rows = [
                    _iam_kpi('MFA Policy',    'NÃO' if no_mfa else 'SIM', not no_mfa),
                    _iam_kpi('Guest Admins',  guest_adm, guest_adm == 0,  ''),
                    _iam_kpi('SP Expirados',  sp_exp,    sp_exp == 0,     ''),
                    _iam_kpi('Owners+',       owners,    owners == 0,     ''),
                    _iam_kpi('Wildcard Roles',wildcard,  wildcard == 0,   ''),
                ]
            else:
                iam_kpi_rows = []
            if iam_kpi_rows:
                iam_kpi_tbl = Table(
                    [row for row in zip(*iam_kpi_rows)],
                    colWidths=[1.25*inch]*5)
                iam_kpi_tbl.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), colors.white),
                    ('LINEAFTER',  (0,0), (3,-1), 0.5, colors.HexColor('#e2e8f0')),
                    ('LINEBEFORE', (0,0), (0,-1), 0.5, colors.HexColor('#e2e8f0')),
                    ('LINEAFTER',  (4,0), (4,-1), 0.5, colors.HexColor('#e2e8f0')),
                    ('LINEABOVE',  (0,0), (-1,0), 3, colors.HexColor('#a855f7')),
                    ('LINEBELOW',  (0,-1), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
                    ('PADDING',    (0,0), (-1,-1), 8),
                    ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
                    ('TOPPADDING', (0,0), (-1,0), 12),
                ]))
                story.append(Paragraph("POSTURA IAM — INDICADORES CHAVE",
                                        ParagraphStyle('IAM_TTL', fontSize=11, fontName='Helvetica-Bold',
                                                       textColor=colors.HexColor('#7c3aed'), spaceAfter=2)))
                story.append(iam_kpi_tbl)
                story.append(Spacer(1, 0.15*inch))

        # Métricas resumidas + ação
        is_iam   = 'IAM' in self.scan_data.get('provider','').upper()
        iam_prov = self.scan_data.get('provider','').upper()
        iam_summ = self.scan_data.get('summary', {}) if is_iam else {}
        target_label = 'Account / Project / Subscription' if is_iam else 'Bucket / Container'
        files_label  = 'Total de Findings' if is_iam else 'Total de Arquivos'
        if iam_prov == 'AWS_IAM':
            size_label = 'Usuários Analisados'
            size_value = str(iam_summ.get('users_total', '-'))
        elif iam_prov == 'GCP_IAM':
            size_label = 'SA Keys Antigas'
            size_value = str(iam_summ.get('sa_keys_old', '-'))
        elif iam_prov == 'AZURE_IAM':
            size_label = 'Guest Admins'
            size_value = str(iam_summ.get('guest_admins', '-'))
        else:
            size_label = 'Tamanho Total'
            size_value = self._format_size(self.size_distribution.get('_total_bytes', 0))
        mt_data = [
            ['Métrica', 'Valor', '% do Total'],
            ['Provider', self.provider_name, ''],
            [target_label, self.scan_data.get('bucket',''), ''],
            [files_label, f"{total:,}", '100%'],
            [size_label, size_value, ''],
            ['Nível de Risco', risk['level'], ''],
        ]
        mt = Table(mt_data, colWidths=[2.8*inch, 2.4*inch, 1.1*inch])
        mt.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,0), self.COLOR_PRIMARY),
            ('TEXTCOLOR',(0,0),(-1,0),  colors.white),
            ('FONTNAME',(0,0),(-1,0),   'Helvetica-Bold'),
            ('FONTSIZE',(0,0),(-1,-1),  10),
            ('PADDING',(0,0),(-1,-1),   9),
            ('GRID',(0,0),(-1,-1),      0.4, colors.HexColor('#e2e8f0')),
            ('ROWBACKGROUNDS',(0,1),(-1,-1), [colors.white, colors.HexColor('#f8fafc')]),
            ('TEXTCOLOR',(1,-1),(2,-1), risk['color']),
            ('FONTNAME',(1,-1),(2,-1),  'Helvetica-Bold'),
        ]))

        # Donut ao lado da tabela de métricas
        if self.ch.get('donut'):
            layout = Table([[
                mt,
                RLImage(self.ch['donut'], width=2.8*inch, height=2.8*inch*0.85),
            ]], colWidths=[3.6*inch, 2.9*inch])
            layout.setStyle(TableStyle([
                ('VALIGN',(0,0),(-1,-1),'TOP'),
                ('ALIGN',(1,0),(1,0),'CENTER'),
            ]))
            story.append(layout)
        else:
            story.append(mt)

        story.append(Spacer(1, 0.12*inch))
        story.append(Paragraph(f"<b>{risk['action']}</b>",
                                ParagraphStyle('ACT', fontSize=11, textColor=risk['color'],
                                               alignment=TA_CENTER, fontName='Helvetica-Bold')))
        story.append(Spacer(1, 0.18*inch))

        # ── Gauge global + parágrafo narrativo ────────────────────────
        narrative = self._generate_narrative()
        narr_para = Paragraph(narrative,
                              ParagraphStyle('NARR', fontSize=10, textColor=colors.HexColor('#1e293b'),
                                             alignment=TA_JUSTIFY, leading=15, spaceAfter=8))
        if self.ch.get('env_gauge'):
            gauge_layout = Table([[
                RLImage(self.ch['env_gauge'], width=2.6*inch, height=2.6*inch*0.68),
                narr_para,
            ]], colWidths=[2.8*inch, 3.8*inch])
            gauge_layout.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN',  (0, 0), (0,  0),  'CENTER'),
                ('LEFTPADDING',  (1, 0), (1, 0), 14),
            ]))
            story.append(gauge_layout)
            story.append(Paragraph("Figura 1 — Risco Global do Ambiente", NOTE))
        else:
            story.append(narr_para)

        story.append(Spacer(1, 0.14*inch))

        # ── Top 10 arquivos mais críticos ─────────────────────────────
        if self.top_critical_files:
            story.append(Paragraph("Arquivos de Maior Risco", H2))
            top_data = [['#', 'Arquivo', 'Severidade', 'Tamanho', 'Motivo']]
            for i, f in enumerate(self.top_critical_files[:10], 1):
                fname = f['file']
                fname_display = ('...' + fname[-42:]) if len(fname) > 45 else fname
                top_data.append([
                    str(i),
                    fname_display,
                    f['severity'],
                    self._format_size(f['size']),
                    (f.get('reason','') or '')[:28],
                ])
            sev_col_map = {'CRITICAL': self.COLOR_CRITICAL, 'HIGH': self.COLOR_HIGH,
                           'MEDIUM': self.COLOR_MEDIUM, 'LOW': self.COLOR_LOW}
            top_t = Table(top_data, colWidths=[0.28*inch, 2.9*inch, 0.75*inch, 0.72*inch, 1.6*inch])
            style_cmds = [
                ('BACKGROUND', (0,0), (-1,0), self.COLOR_PRIMARY),
                ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
                ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE',   (0,0), (-1,-1), 8.5),
                ('PADDING',    (0,0), (-1,-1), 6),
                ('GRID',       (0,0), (-1,-1), 0.4, colors.HexColor('#e5e7eb')),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, self.COLOR_BG_LIGHT]),
                ('ALIGN',      (0,0), (0,-1), 'CENTER'),
                ('ALIGN',      (2,0), (3,-1), 'CENTER'),
            ]
            for row_i, f in enumerate(self.top_critical_files[:10], 1):
                col = sev_col_map.get(f['severity'], self.COLOR_LOW)
                style_cmds.append(('TEXTCOLOR', (2, row_i), (2, row_i), col))
                style_cmds.append(('FONTNAME',  (2, row_i), (2, row_i), 'Helvetica-Bold'))
            top_t.setStyle(TableStyle(style_cmds))
            story.append(KeepTogether([top_t, Spacer(1, 0.08*inch)]))

        story.append(PageBreak())

        # ── 2. DISTRIBUIÇÃO DE CRITICIDADE ────────────────────────────
        sec2_title = "2. DISTRIBUIÇÃO DE FINDINGS IAM" if is_iam else "2. DISTRIBUIÇÃO DE CRITICIDADE"
        sec2_desc  = ("Distribuição dos findings por nível de severidade e categoria de check IAM."
                      if is_iam else
                      "Distribuição dos arquivos por nível de severidade e tipo de conteúdo exposto.")
        story.append(Paragraph(sec2_title, H1))
        story.append(Paragraph(sec2_desc, styles['Normal']))
        story.append(Spacer(1, 0.1*inch))

        # ── Funil + Top Achados por Categoria (side by side) ─────────
        if self.ch.get('funnel'):
            _CAT_PDF = [
                ('Credenciais',   lambda k: any(x in k for x in ['password','credential','secret','token','api_key'])),
                ('Chaves Crypto', lambda k: any(x in k for x in ['.pem','.key','.crt','.p12','.pfx','rsa'])),
                ('DB / Backups',  lambda k: any(x in k for x in ['.sql','.db','.sqlite','backup','.bak','.dump','.gz','.zip','.tar'])),
                ('Código-fonte',  lambda k: any(x in k for x in ['.py','.js','.java','.php','.rb','.go'])),
                ('Configs',       lambda k: any(x in k for x in ['.env','.yaml','.yml','.conf','.ini','config'])),
                ('Outros',        lambda k: True),
            ]
            _grps = {}
            for _f in self.scan_data.get('files', []):
                _n  = _f.get('key', _f.get('name', '')).lower()
                _sv = _f.get('severity', 'LOW').lower()
                for _cn, _rule in _CAT_PDF:
                    if _rule(_n):
                        if _cn not in _grps:
                            _grps[_cn] = {'critical':0,'high':0,'medium':0,'low':0,'total':0}
                        if _sv in _grps[_cn]: _grps[_cn][_sv] += 1
                        _grps[_cn]['total'] += 1
                        break
            _cat_list = []
            for _cn, _g in _grps.items():
                if _g['total'] == 0: continue
                _raw = (_g['critical']*10 + _g['high']*7 + _g['medium']*4 + _g['low']*1) / _g['total']
                _cat_list.append((_cn, _g, min(10.0, _raw)))
            _cat_list.sort(key=lambda x: (-x[2], -x[1]['total']))
            _cat_list = _cat_list[:6]

            if _cat_list:
                _max_tot = max(c[1]['total'] for c in _cat_list)
                _SH = ParagraphStyle('_SH', fontSize=7, fontName='Helvetica-Bold',
                                     textColor=colors.HexColor('#64748b'), alignment=TA_CENTER)
                _ta = [[Paragraph('<b>#</b>', _SH),
                        Paragraph('<b>Categoria</b>', _SH),
                        Paragraph('<b>Score</b>', _SH),
                        Paragraph('<b>Arquivos</b>', _SH),
                        Paragraph('', _SH)]]
                for _i, (_cn, _g, _score) in enumerate(_cat_list, 1):
                    _sc = (self.COLOR_CRITICAL if _score >= 7 else
                           self.COLOR_HIGH     if _score >= 4 else
                           self.COLOR_MEDIUM   if _score >= 2 else self.COLOR_LOW)
                    _filled = max(0.06, _g['total'] / _max_tot)
                    _BW = 1.1 * inch
                    _bt = Table([['', '']], colWidths=[_filled * _BW, max(1, (1 - _filled) * _BW)],
                                rowHeights=[10])
                    _bt.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, 0), _sc),
                        ('BACKGROUND', (1, 0), (1, 0), colors.HexColor('#e2e8f0')),
                        ('TOPPADDING',    (0, 0), (-1, -1), 0),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                        ('LEFTPADDING',   (0, 0), (-1, -1), 0),
                        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
                    ]))
                    _ta.append([
                        Paragraph(f'<b>{_i}</b>', ParagraphStyle('_idx', fontSize=8,
                            fontName='Helvetica-Bold', textColor=colors.HexColor('#6b7280'),
                            alignment=TA_CENTER)),
                        Paragraph(f'<b>{_cn}</b>', ParagraphStyle('_cn', fontSize=8,
                            fontName='Helvetica-Bold', textColor=colors.HexColor('#1e293b'))),
                        Paragraph(f'<b>{_score:.1f}</b>', ParagraphStyle('_sc', fontSize=9,
                            fontName='Helvetica-Bold', textColor=_sc, alignment=TA_CENTER)),
                        _bt,
                        Paragraph(f'<b>{_g["total"]:,}</b>', ParagraphStyle('_ct', fontSize=8,
                            fontName='Helvetica-Bold', textColor=colors.HexColor('#374151'),
                            alignment=TA_RIGHT)),
                    ])
                _ta_tbl = Table(_ta, colWidths=[0.2*inch, 1.1*inch, 0.45*inch, 1.1*inch, 0.42*inch])
                _ta_tbl.setStyle(TableStyle([
                    ('BACKGROUND',    (0, 0), (-1, 0), colors.HexColor('#f8fafc')),
                    ('LINEABOVE',     (0, 0), (-1, 0), 1.5, self.COLOR_PRIMARY),
                    ('ROWBACKGROUNDS',(0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
                    ('GRID',          (0, 0), (-1, -1), 0.4, colors.HexColor('#e2e8f0')),
                    ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING',    (0, 0), (-1, -1), 7),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
                    ('LEFTPADDING',   (0, 0), (-1, -1), 5),
                    ('RIGHTPADDING',  (0, 0), (-1, -1), 5),
                ]))
                _n_sev  = sum(1 for k in ['critical','high','medium','low'] if sd.get(k, 0) > 0)
                _f_h    = max(0.9, _n_sev * 0.52) * inch
                _f_img  = RLImage(self.ch['funnel'], width=3.4*inch, height=_f_h)
                _side   = Table([[_f_img, _ta_tbl]], colWidths=[3.55*inch, 3.42*inch])
                _side.setStyle(TableStyle([
                    ('VALIGN',       (0, 0), (-1, -1), 'TOP'),
                    ('RIGHTPADDING', (0, 0), (0, 0),  12),
                    ('LEFTPADDING',  (1, 0), (1, 0),   0),
                ]))
                story.append(_side)
                story.append(Paragraph(
                    "Figura 2 — Achados por Severidade (esq.)  |  Top Categorias de Risco (dir.)", NOTE))
                story.append(Spacer(1, 0.18*inch))

        if self.ch.get('donut') or self.ch.get('type_pie'):
            left_img  = RLImage(self.ch['donut'],    width=3.2*inch, height=3.2*inch*0.78) if self.ch.get('donut') else Paragraph('', styles['Normal'])
            right_img = RLImage(self.ch['type_pie'], width=3.2*inch, height=3.2*inch*0.78) if self.ch.get('type_pie') else Paragraph('', styles['Normal'])
            row = Table([[left_img, right_img]], colWidths=[3.4*inch, 3.4*inch])
            row.setStyle(TableStyle([('ALIGN',(0,0),(-1,-1),'CENTER'),
                                     ('VALIGN',(0,0),(-1,-1),'MIDDLE')]))
            story.append(row)
            story.append(Paragraph("Figura 3 — Severidade (esq.)  |  Tipo de Arquivo (dir.)", NOTE))
        # Altura proporcional ao nº de categorias para evitar whitespace excessivo
        def _cat_of(key):
            k = key.lower()
            if any(x in k for x in ['.sql','.db','.sqlite','backup','.bak','.dump','.gz','.zip','.tar']): return 'DB'
            if any(x in k for x in ['password','credential','secret','token','api_key']): return 'CRED'
            if any(x in k for x in ['.pem','.key','.crt','.p12','.pfx','rsa']): return 'KEY'
            if any(x in k for x in ['.py','.js','.java','.php','.rb','.go']): return 'CODE'
            if any(x in k for x in ['.env','.yaml','.yml','.conf','.ini','config']): return 'CFG'
            return 'OTHER'
        _n_cats     = max(1, len({_cat_of(f.get('key', f.get('name', ''))) for f in self.scan_data.get('files', [])}))
        _cat_height = max(0.9*inch, _n_cats * 0.38*inch + 0.55*inch)
        add_chart(self.ch.get('cat_bars'), width=5.8*inch, height=_cat_height,
                  caption="Figura 4 — Distribuição por Categoria de Risco")

        story.append(PageBreak())

        # ── 3. SCORE DE CRITICIDADE POR FINDING / ARQUIVO ─────────────
        sec3_title = "3. SCORE DE CRITICIDADE POR FINDING IAM" if is_iam else "3. SCORE DE CRITICIDADE POR ARQUIVO"
        sec3_desc  = ("Score calculado por finding: base de severidade. "
                      "Zonas coloridas indicam faixas de risco IAM." if is_iam else
                      "Score calculado por arquivo: base de severidade + bônus proporcional ao tamanho. "
                      "Zonas coloridas indicam faixas de risco.")
        story.append(Paragraph(sec3_title, H1))
        story.append(Paragraph(sec3_desc, styles['Normal']))
        add_chart(self.ch.get('risk_all'), width=5.8*inch,
                  caption="Figura 5 — Score de Criticidade por Finding" if is_iam else "Figura 5 — Score de Criticidade por Arquivo")
        if not is_iam:
            story.append(Paragraph("Volume de dados expostos por arquivo:", H2))
            add_chart(self.ch.get('size'), width=5.8*inch,
                      caption="Figura 6 — Volume de Dados Expostos por Arquivo")
        story.append(PageBreak())

        # ── 4. MAPA DE CALOR / IAM CHECK BARS ────────────────────────
        if is_iam:
            if self.ch.get('cat_bars'):
                story.append(Paragraph("4. DISTRIBUIÇÃO DE FINDINGS POR CATEGORIA IAM", H1))
                story.append(Paragraph(
                    "Findings agrupados por categoria de check: MFA, access keys, permissões excessivas, "
                    "cross-account trust, configuração de conta e inatividade.",
                    styles['Normal']))
                add_chart(self.ch.get('cat_bars'), width=5.8*inch,
                          caption="Figura 6 — Findings IAM por Categoria e Severidade")
        else:
            story.append(Paragraph("4. MAPA DE CALOR DE RISCO MULTIDIMENSIONAL", H1))
            story.append(Paragraph(
                "Cada arquivo é avaliado em 5 dimensões: Exposição Pública, Volume de Dados, "
                "Tipo de Arquivo, Risco de Privacidade e Score Final. Cores quentes = maior risco.",
                styles['Normal']))
            add_chart(self.ch.get('heatmap'), width=6.0*inch,
                      caption="Figura 6 — Mapa de Calor Multidimensional de Risco")
        story.append(PageBreak())

        # ── 5. GAUGES INDIVIDUAIS ──────────────────────────────────────
        if self.ch.get('gauges'):
            story.append(Paragraph("5. ANÁLISE INDIVIDUAL POR ARQUIVO (GAUGES)", H1))
            story.append(Paragraph(
                "Velocímetro de criticidade individual para cada arquivo vulnerável. "
                "O ponteiro indica o score calculado (0–100).", styles['Normal']))
            story.append(Spacer(1, 0.15*inch))

            # Grade 3 colunas de gauges
            COLS = 3
            gauges = self.ch['gauges']
            files  = self.scan_data.get('files', [])[:len(gauges)]
            for row_start in range(0, len(gauges), COLS):
                chunk = gauges[row_start:row_start+COLS]
                imgs  = [RLImage(p, width=1.95*inch, height=1.95*inch*0.62) for p in chunk]
                while len(imgs) < COLS:
                    imgs.append(Paragraph('', styles['Normal']))
                row_t = Table([imgs], colWidths=[2.1*inch]*COLS)
                row_t.setStyle(TableStyle([
                    ('ALIGN',(0,0),(-1,-1),'CENTER'),
                    ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                    ('BOTTOMPADDING',(0,0),(-1,-1),4),
                ]))
                story.append(row_t)
            story.append(Paragraph(
                f"Figura 7 — Gauges de criticidade individual ({len(gauges)} arquivos)", NOTE))
            story.append(PageBreak())
            next_sec = 6
        else:
            next_sec = 5

        # ── VULNERABILIDADES ──────────────────────────────────────────
        if self.vulnerabilities:
            story.append(Paragraph(f"{next_sec}. VULNERABILIDADES IDENTIFICADAS", H1))
            for vuln in self.vulnerabilities:
                story.append(Paragraph(
                    f"{vuln['icon']} {vuln['name']} ({vuln['count']} arquivos)", H2))
                vt = Table(
                    [['Arquivo','Sev.','Tamanho','Motivo']] +
                    [[item['file'][:45]+'...' if len(item['file'])>45 else item['file'],
                      item['severity'],
                      self._format_size(item['size']),
                      item.get('reason','')[:30]] for item in vuln['items'][:10]],
                    colWidths=[2.8*inch, 0.7*inch, 0.8*inch, 1.8*inch])
                vt.setStyle(TableStyle([
                    ('BACKGROUND',(0,0),(-1,0),self.COLOR_SECONDARY),
                    ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                    ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
                    ('FONTSIZE',(0,0),(-1,-1),8),
                    ('PADDING',(0,0),(-1,-1),6),
                    ('GRID',(0,0),(-1,-1),0.5,colors.grey),
                    ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,self.COLOR_BG_LIGHT]),
                ]))
                story.append(vt)
                story.append(Spacer(1, 0.15*inch))
            story.append(PageBreak())
            next_sec += 1

        # ── COMPARATIVO ───────────────────────────────────────────────
        comp = self._build_comparison()
        if comp:
            story.append(Paragraph(f"{next_sec}. COMPARATIVO COM SCAN ANTERIOR", H1))
            story.append(Paragraph(
                f"Scan anterior realizado em: <b>{comp['prev_timestamp']}</b>",
                styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            ct_data = [['Métrica','Anterior','Atual','Variação']]
            for lbl, prev_val, curr_val, delta in [
                ('Total', comp['prev_total'], comp['curr_total'],
                 comp['curr_total']-comp['prev_total']),
                ('CRÍTICO',
                 self.previous_scan['severity_distribution'].get('critical',0),
                 self.scan_data['severity_distribution'].get('critical',0),
                 comp['delta_critical']),
                ('ALTO',
                 self.previous_scan['severity_distribution'].get('high',0),
                 self.scan_data['severity_distribution'].get('high',0),
                 comp['delta_high']),
                ('Risk Score', self.previous_scan.get('risk_score',0),
                 self.scan_data.get('risk_score',0), comp['delta_score']),
            ]:
                ct_data.append([lbl, str(prev_val), str(curr_val), self._delta_str(delta)])

            comp_t = Table(ct_data, colWidths=[2.5*inch,1.2*inch,1.2*inch,1.2*inch])
            comp_t.setStyle(TableStyle([
                ('BACKGROUND',(0,0),(-1,0),self.COLOR_PRIMARY),
                ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
                ('FONTSIZE',(0,0),(-1,-1),10),
                ('PADDING',(0,0),(-1,-1),10),
                ('GRID',(0,0),(-1,-1),0.5,self.COLOR_PRIMARY),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,self.COLOR_BG_LIGHT]),
                ('ALIGN',(1,0),(-1,-1),'CENTER'),
            ]))
            story.append(comp_t)
            story.append(PageBreak())
            next_sec += 1

        # ── RECOMENDAÇÕES ─────────────────────────────────────────────
        story.append(Paragraph(f"{next_sec}. RECOMENDAÇÕES PRIORITÁRIAS", H1))
        effort_map = {'CRÍTICA':'Baixo','ALTA':'Médio','MÉDIA':'Alto','BAIXA':'Baixo'}
        impact_map = {'CRÍTICA':'Máximo','ALTA':'Alto','MÉDIA':'Médio','BAIXA':'Baixo'}
        pcols = {'CRÍTICA':self.COLOR_CRITICAL,'ALTA':self.COLOR_HIGH,
                 'MÉDIA':self.COLOR_MEDIUM,'BAIXA':self.COLOR_LOW}
        for rec in self.recommendations:
            story.append(Paragraph(
                f"[{rec['priority']}] {rec['title']}",
                ParagraphStyle('RT', fontSize=11, fontName='Helvetica-Bold',
                               textColor=pcols.get(rec['priority'],self.COLOR_MEDIUM), spaceAfter=5)))
            story.append(Paragraph(rec['description'], styles['Normal']))
            meta = Table([[
                'Prazo', rec.get('timeline','A definir'),
                'Responsável', rec.get('responsible','A definir'),
                'Esforço', effort_map.get(rec['priority'],'Médio'),
                'Impacto', impact_map.get(rec['priority'],'Alto'),
            ]], colWidths=[0.8*inch,1.3*inch,1.0*inch,1.7*inch,0.8*inch,0.9*inch,0.7*inch,0.8*inch])
            meta.setStyle(TableStyle([
                ('BACKGROUND',(0,0),(-1,-1),self.COLOR_BG_LIGHT),
                ('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),
                ('FONTNAME',(2,0),(2,-1),'Helvetica-Bold'),
                ('FONTNAME',(4,0),(4,-1),'Helvetica-Bold'),
                ('FONTNAME',(6,0),(6,-1),'Helvetica-Bold'),
                ('FONTSIZE',(0,0),(-1,-1),8),
                ('PADDING',(0,0),(-1,-1),5),
                ('GRID',(0,0),(-1,-1),0.3,colors.grey),
                ('ALIGN',(0,0),(-1,-1),'CENTER'),
            ]))
            story.append(Spacer(1,0.05*inch))
            story.append(meta)
            story.append(Spacer(1,0.05*inch))
            story.append(Paragraph("<b>Ações:</b>", styles['Normal']))
            for a in rec['actions'][:5]:
                story.append(Paragraph(f"• {a}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        story.append(PageBreak())
        next_sec += 1

        # ── COMPLIANCE ────────────────────────────────────────────────
        if self.compliance_status:
            story.append(Paragraph(f"{next_sec}. STATUS DE CONFORMIDADE", H1))
            comp_t2 = Table(
                [['Framework','Status','Observações']] +
                [[fw['name'],fw['status'],fw['issues']] for fw in self.compliance_status],
                colWidths=[1.5*inch,1.8*inch,2.8*inch])
            comp_t2.setStyle(TableStyle([
                ('BACKGROUND',(0,0),(-1,0),self.COLOR_PRIMARY),
                ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
                ('FONTSIZE',(0,0),(-1,-1),9),
                ('PADDING',(0,0),(-1,-1),8),
                ('GRID',(0,0),(-1,-1),0.5,self.COLOR_PRIMARY),
                ('VALIGN',(0,0),(-1,-1),'TOP'),
            ]))
            story.append(comp_t2)
            story.append(Spacer(1,0.3*inch))
            next_sec += 1

        # ── CONCLUSÃO ─────────────────────────────────────────────────
        story.append(Paragraph(f"{next_sec}. CONCLUSÕES E PRÓXIMOS PASSOS", H1))
        story.append(Paragraph(
            f"Este relatório identificou <b>{risk['critical_count']} vulnerabilidades críticas</b> e "
            f"<b>{risk['high_count']} de alto risco</b> no ambiente {self.provider_name}. "
            "É fundamental executar as ações prioritárias nos prazos estabelecidos.<br/><br/>"
            "<b>Próximos Passos:</b><br/>"
            "1. Reunião de alinhamento com stakeholders (48h)<br/>"
            "2. Execução das ações críticas (0-7 dias)<br/>"
            "3. Controles preventivos (30-60 dias)<br/>"
            "4. Re-auditoria de segurança (90 dias)<br/>"
            "5. Programa de segurança contínua",
            styles['Normal']))
        story.append(Spacer(1,0.5*inch))
        story.append(Paragraph(
            f"<b>Security Multicloud Scanner</b><br/>{self.provider_name} | "
            "Relatório gerado automaticamente<br/>Este documento contém informações confidenciais",
            ParagraphStyle('SIG', fontSize=8, textColor=colors.grey, alignment=TA_CENTER)))

        print("🔨 Construindo PDF...")
        doc.build(story, onFirstPage=self._create_cover_page,
                  onLaterPages=self._create_header_footer)

        report_hash = self._compute_hash(str(filepath))
        (filepath.with_suffix('.sha256')).write_text(
            f"SHA-256: {report_hash}\nArquivo: {filename}\nGerado em: {self.report_date.isoformat()}\n")
        print(f"✅ PDF: {filepath}")
        return str(filepath)

    # ══════════════════════════════════════════════════════════════════
    # DOCX GENERATION
    # ══════════════════════════════════════════════════════════════════
    def _docx_set_bg(self, cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color.replace('#',''))
        tcPr.append(shd)

    def _docx_remove_borders(self, cell):
        tcPr = cell._tc.get_or_add_tcPr()
        bdr  = OxmlElement('w:tcBorders')
        for side in ['top','bottom','left','right','insideH','insideV']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            bdr.append(el)
        tcPr.append(bdr)

    def _docx_heading(self, doc, text, level=1):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt({1:16,2:13,3:11}.get(level,11))
        run.font.color.rgb = RGBColor(0x1e,0x3a,0x8a) if level==1 else RGBColor(0x3b,0x82,0xf6)
        run.font.name = 'Arial'
        p.paragraph_format.space_before = Pt(14 if level==1 else 10)
        p.paragraph_format.space_after  = Pt(6)
        if level == 1:
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
            bot.set(qn('w:space'),'4');   bot.set(qn('w:color'),'3b82f6')
            pBdr.append(bot); pPr.append(pBdr)
        return p

    def _docx_body(self, doc, text, color='333333', bold=False, size=10, after=4):
        p   = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(after)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(text)
        run.font.size  = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name  = 'Arial'
        run.bold       = bold
        return p

    def _docx_add_chart(self, doc, path, width_in=6.0, caption=None):
        if not path or not Path(path).exists(): return
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        p.add_run().add_picture(path, width=Inches(width_in))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(8)
            r = cap.add_run(caption)
            r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x6b,0x72,0x80)
            r.font.name = 'Arial'; r.italic = True

    def generate_docx(self) -> str:
        if not DOCX_AVAILABLE:
            print("⚠️  python-docx indisponível"); return None
        print("📝 Gerando DOCX...")
        bucket   = self.scan_data.get('bucket','scan').replace('.','_').replace('/','_')
        filename = f"relatorio_{bucket}_{self.timestamp}.docx"
        filepath = self.output_dir / filename

        doc = Document()
        for sec in doc.sections:
            sec.page_height   = Inches(11.69); sec.page_width    = Inches(8.27)
            sec.left_margin   = Inches(0.85);  sec.right_margin  = Inches(0.85)
            sec.top_margin    = Inches(0.9);   sec.bottom_margin = Inches(0.9)

        sd       = self.scan_data.get('severity_distribution', {})
        total    = len(self.scan_data.get('files', []))
        risk     = self.risk_level
        is_iam   = 'IAM' in self.scan_data.get('provider','').upper()
        iam_summ = self.scan_data.get('summary', {}) if is_iam else {}

        # ── Capa ──────────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(60)
        capa_title = (f'RELATÓRIO DE SEGURANÇA EXECUTIVO\nCSPM — {self.provider_name}' if is_iam
                      else 'RELATÓRIO DE SEGURANÇA EXECUTIVO\nCOM ANÁLISE DE CRITICIDADE')
        r = p.add_run(capa_title)
        r.font.size = Pt(24); r.bold = True
        r.font.color.rgb = RGBColor(0x1e,0x3a,0x8a); r.font.name = 'Arial'
        subtitle = (f'Auditoria de Postura IAM — {self.provider_name}' if is_iam
                    else f'Auditoria de Storage Multicloud — {self.provider_name}')
        self._docx_body(doc, subtitle, color='6b7280', size=13).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        if self.client_info.get('name'):
            tbl = doc.add_table(rows=4, cols=2)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            rows_data = [('Cliente', self.client_info.get('name','')),
                         ('Contato', self.client_info.get('contact','')),
                         ('Data', self.report_date.strftime('%d/%m/%Y às %H:%M')),
                         ('Alvo', self.scan_data.get('bucket',''))]
            for i,(k,v) in enumerate(rows_data):
                self._docx_set_bg(tbl.rows[i].cells[0], 'e8eef8')
                tbl.rows[i].cells[0].width = Inches(1.4)
                tbl.rows[i].cells[1].width = Inches(3.0)
                r0 = tbl.rows[i].cells[0].paragraphs[0].add_run(k)
                r0.bold=True; r0.font.size=Pt(10); r0.font.name='Arial'
                r0.font.color.rgb=RGBColor(0x1e,0x3a,0x8a)
                r1 = tbl.rows[i].cells[1].paragraphs[0].add_run(v)
                r1.font.size=Pt(10); r1.font.name='Arial'

        doc.add_page_break()

        # ── 1. Resumo ─────────────────────────────────────────────────
        self._docx_heading(doc, '1. RESUMO EXECUTIVO')
        self._docx_body(doc, f"Nível de Risco: {risk['level']} ({risk['score']}%) — {risk['action']}",
                         bold=True, color='dc2626' if risk['level']=='CRÍTICO' else '1e3a8a')
        doc.add_paragraph()

        tbl = doc.add_table(rows=2, cols=4)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        if is_iam:
            iam_prov_docx = self.scan_data.get('provider','').upper()
            if iam_prov_docx == 'GCP_IAM':
                extra_lbl, extra_val = 'SA Keys Antigas', str(iam_summ.get('sa_keys_old', 0))
            elif iam_prov_docx == 'AZURE_IAM':
                extra_lbl, extra_val = 'Guest Admins', str(iam_summ.get('guest_admins', 0))
            else:
                extra_lbl, extra_val = 'Sem MFA', str(iam_summ.get('users_no_mfa', 0))
            metrics = [('Findings', str(total), '1e3a8a'),
                       ('CRÍTICO',  str(sd.get('critical',0)), 'dc2626'),
                       ('ALTO',     str(sd.get('high',0)),     'ea580c'),
                       (extra_lbl,  extra_val, '7c3aed')]
        else:
            metrics = [('Arquivos', str(total), '1e3a8a'),
                       ('CRÍTICO', str(sd.get('critical',0)), 'dc2626'),
                       ('ALTO',    str(sd.get('high',0)),    'ea580c'),
                       ('Tamanho',  self._format_size(self.size_distribution.get('_total_bytes',0)), '1e3a8a')]
        for i,(lbl,val,col) in enumerate(metrics):
            self._docx_set_bg(tbl.rows[0].cells[i], '1e3a8a')
            r0 = tbl.rows[0].cells[i].paragraphs[0].add_run(lbl)
            r0.bold=True; r0.font.size=Pt(9); r0.font.name='Arial'
            r0.font.color.rgb=RGBColor(255,255,255)
            tbl.rows[0].cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            r1 = tbl.rows[1].cells[i].paragraphs[0].add_run(val)
            r1.bold=True; r1.font.size=Pt(16); r1.font.name='Arial'
            r1.font.color.rgb=RGBColor.from_string(col)
            tbl.rows[1].cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # ── 2. Distribuição de criticidade ───────────────────────────
        if is_iam:
            self._docx_heading(doc, '2. DISTRIBUIÇÃO DE FINDINGS IAM')
            self._docx_body(doc, 'Distribuição dos findings por severidade e categoria de check IAM.')
            self._docx_add_chart(doc, self.ch.get('donut'), width_in=5.5,
                                  caption='Figura 1 — Distribuição de Severidade dos Findings')
            if self.ch.get('cat_bars'):
                self._docx_add_chart(doc, self.ch['cat_bars'], width_in=6.0,
                                      caption='Figura 2 — Findings IAM por Categoria')
        else:
            self._docx_heading(doc, '2. DISTRIBUIÇÃO DE CRITICIDADE')
            self._docx_body(doc, 'Gráfico de distribuição dos arquivos por nível de severidade e tipo de conteúdo.')
            self._docx_add_chart(doc, self.ch.get('donut'), width_in=5.5,
                                  caption='Figura 1 — Distribuição de Severidade')
            if self.ch.get('type_pie'):
                self._docx_add_chart(doc, self.ch['type_pie'], width_in=4.5,
                                      caption='Figura 2 — Distribuição por Tipo de Arquivo')
        doc.add_page_break()

        # ── 3. Score por finding / arquivo ───────────────────────────
        if is_iam:
            self._docx_heading(doc, '3. SCORE DE CRITICIDADE POR FINDING IAM')
            self._docx_body(doc, 'Score calculado por finding com base na severidade do check IAM.')
            self._docx_add_chart(doc, self.ch.get('risk_all'), width_in=6.2,
                                  caption='Figura 3 — Score de Criticidade por Finding IAM')
        else:
            self._docx_heading(doc, '3. SCORE DE CRITICIDADE POR ARQUIVO')
            self._docx_body(doc, 'Score calculado: base de severidade + bônus proporcional ao tamanho do arquivo.')
            self._docx_add_chart(doc, self.ch.get('risk_all'), width_in=6.2,
                                  caption='Figura 3 — Score de Criticidade por Arquivo')
            self._docx_add_chart(doc, self.ch.get('size'), width_in=6.2,
                                  caption='Figura 4 — Volume de Dados Expostos por Arquivo')
        doc.add_page_break()

        # ── 4. Mapa de calor ─────────────────────────────────────────
        if not is_iam:
            self._docx_heading(doc, '4. MAPA DE CALOR DE RISCO MULTIDIMENSIONAL')
            self._docx_body(doc,
                'Avalia cada arquivo em 5 dimensões: Exposição Pública, Volume, Tipo, '
                'Risco de Privacidade e Score Final. Cores quentes = maior risco.')
            self._docx_add_chart(doc, self.ch.get('heatmap'), width_in=6.3,
                                  caption='Figura 5 — Mapa de Calor Multidimensional')
            doc.add_page_break()

        # ── 5. Gauges individuais ─────────────────────────────────────
        gauges = self.ch.get('gauges', [])
        if gauges:
            self._docx_heading(doc, '5. ANÁLISE INDIVIDUAL POR ARQUIVO (GAUGES)')
            self._docx_body(doc,
                'Velocímetro de criticidade individual para cada arquivo vulnerável identificado no scan.')
            doc.add_paragraph()

            files_list = self.scan_data.get('files', [])[:len(gauges)]
            max_size   = max((float(f.get('size',0)) for f in files_list), default=1) or 1

            COLS = 3
            for row_start in range(0, len(gauges), COLS):
                chunk = gauges[row_start:row_start+COLS]
                fchunk = files_list[row_start:row_start+COLS]

                # Linha de imagens
                img_tbl = doc.add_table(rows=2, cols=COLS)
                img_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for col_i, (gpath, f) in enumerate(zip(chunk, fchunk)):
                    # Imagem
                    ic = img_tbl.rows[0].cells[col_i]
                    ic.width = Inches(2.1)
                    self._docx_remove_borders(ic)
                    pi = ic.paragraphs[0]
                    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if Path(gpath).exists():
                        pi.add_run().add_picture(gpath, width=Inches(2.0))
                    # Caption
                    cc = img_tbl.rows[1].cells[col_i]
                    cc.width = Inches(2.1)
                    self._docx_remove_borders(cc)
                    score = min(100, self.SEV_BASE.get(f.get('severity','LOW'),18) +
                                (float(f.get('size',0)) / max_size) * 12)
                    sev   = f.get('severity','LOW')
                    name  = f.get('key', f.get('name',''))[-20:]
                    pc = cc.paragraphs[0]
                    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rc = pc.add_run(f"{name}\n{sev} | {score:.0f} pts")
                    rc.font.size = Pt(7.5); rc.font.name = 'Arial'
                    rc.font.color.rgb = RGBColor(0x6b,0x72,0x80)

                # Preencher células vazias se necessário
                for col_i in range(len(chunk), COLS):
                    for row_i in range(2):
                        self._docx_remove_borders(img_tbl.rows[row_i].cells[col_i])

            doc.add_paragraph()
            doc.add_page_break()
            next_sec = 6
        else:
            next_sec = 5

        # ── Vulnerabilidades ──────────────────────────────────────────
        if self.vulnerabilities:
            self._docx_heading(doc, f'{next_sec}. VULNERABILIDADES IDENTIFICADAS')
            for vuln in self.vulnerabilities:
                self._docx_heading(doc, f"{vuln['icon']} {vuln['name']} ({vuln['count']} arquivos)", level=2)
                vt = doc.add_table(rows=1+len(vuln['items'][:10]), cols=4)
                vt.style = 'Table Grid'
                for i, h in enumerate(['Arquivo','Severidade','Tamanho','Motivo']):
                    self._docx_set_bg(vt.rows[0].cells[i], '3b82f6')
                    r = vt.rows[0].cells[i].paragraphs[0].add_run(h)
                    r.bold=True; r.font.size=Pt(9); r.font.name='Arial'
                    r.font.color.rgb=RGBColor(255,255,255)
                for ri, item in enumerate(vuln['items'][:10], 1):
                    vals = [item['file'][-50:], item['severity'],
                            self._format_size(item['size']), item.get('reason','')[:35]]
                    for ci, v in enumerate(vals):
                        c = vt.rows[ri].cells[ci]
                        r = c.paragraphs[0].add_run(v)
                        r.font.size=Pt(8.5); r.font.name='Arial'
                doc.add_paragraph()
            next_sec += 1

        # ── Comparativo ───────────────────────────────────────────────
        comp = self._build_comparison()
        if comp:
            self._docx_heading(doc, f'{next_sec}. COMPARATIVO COM SCAN ANTERIOR')
            self._docx_body(doc, f"Scan anterior: {comp['prev_timestamp']}")
            ct = doc.add_table(rows=5, cols=4)
            ct.style = 'Table Grid'
            for i, h in enumerate(['Métrica','Anterior','Atual','Variação']):
                self._docx_set_bg(ct.rows[0].cells[i], '1e3a8a')
                r = ct.rows[0].cells[i].paragraphs[0].add_run(h)
                r.bold=True; r.font.size=Pt(9); r.font.name='Arial'
                r.font.color.rgb=RGBColor(255,255,255)
            comp_rows = [
                ('Total', comp['prev_total'], comp['curr_total'],
                 comp['curr_total']-comp['prev_total']),
                ('CRÍTICO',
                 self.previous_scan['severity_distribution'].get('critical',0),
                 sd.get('critical',0), comp['delta_critical']),
                ('ALTO',
                 self.previous_scan['severity_distribution'].get('high',0),
                 sd.get('high',0), comp['delta_high']),
                ('Risk Score', self.previous_scan.get('risk_score',0),
                 self.scan_data.get('risk_score',0), comp['delta_score']),
            ]
            for ri,(lbl,prev,curr,delta) in enumerate(comp_rows, 1):
                vals = [lbl, str(prev), str(curr), self._delta_str(delta)]
                for ci,v in enumerate(vals):
                    ct.rows[ri].cells[ci].paragraphs[0].add_run(v).font.size = Pt(9)
            next_sec += 1

        # ── Recomendações ─────────────────────────────────────────────
        self._docx_heading(doc, f'{next_sec}. RECOMENDAÇÕES PRIORITÁRIAS')
        effort_map = {'CRÍTICA':'Baixo','ALTA':'Médio','MÉDIA':'Alto','BAIXA':'Baixo'}
        impact_map = {'CRÍTICA':'Máximo','ALTA':'Alto','MÉDIA':'Médio','BAIXA':'Baixo'}
        prio_hex   = {'CRÍTICA':'dc2626','ALTA':'ea580c','MÉDIA':'ca8a04','BAIXA':'16a34a'}
        for rec in self.recommendations:
            # Cabeçalho colorido
            hdr = doc.add_table(rows=1, cols=2)
            hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._docx_set_bg(hdr.rows[0].cells[0], prio_hex.get(rec['priority'],'3b82f6'))
            self._docx_set_bg(hdr.rows[0].cells[1], prio_hex.get(rec['priority'],'3b82f6'))
            hdr.rows[0].cells[0].width = Inches(4.5)
            hdr.rows[0].cells[1].width = Inches(1.9)
            r0 = hdr.rows[0].cells[0].paragraphs[0].add_run(f"[{rec['priority']}]  {rec['title']}")
            r0.bold=True; r0.font.size=Pt(10.5); r0.font.name='Arial'
            r0.font.color.rgb=RGBColor(255,255,255)
            r1 = hdr.rows[0].cells[1].paragraphs[0].add_run(f"Prazo: {rec.get('timeline','')}")
            r1.font.size=Pt(9); r1.font.name='Arial'
            r1.font.color.rgb=RGBColor(255,255,255)
            hdr.rows[0].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

            # Meta
            meta = doc.add_table(rows=1, cols=4)
            meta.style = 'Table Grid'
            for ci,(lbl,val) in enumerate([
                ('Responsável', rec.get('responsible','')),
                ('Esforço',     effort_map.get(rec['priority'],'')),
                ('Impacto',     impact_map.get(rec['priority'],'')),
                ('Prazo',       rec.get('timeline','')),
            ]):
                self._docx_set_bg(meta.rows[0].cells[ci], 'f8fafc')
                p = meta.rows[0].cells[ci].paragraphs[0]
                rl = p.add_run(f"{lbl}: ")
                rl.bold=True; rl.font.size=Pt(9); rl.font.name='Arial'
                rv = p.add_run(val)
                rv.font.size=Pt(9); rv.font.name='Arial'

            # Descrição e ações
            self._docx_body(doc, rec['description'], after=2)
            for a in rec['actions'][:5]:
                bp = doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)
                br = bp.add_run(a)
                br.font.size=Pt(9.5); br.font.name='Arial'
            doc.add_paragraph()
        next_sec += 1

        # ── Compliance ────────────────────────────────────────────────
        if self.compliance_status:
            self._docx_heading(doc, f'{next_sec}. STATUS DE CONFORMIDADE')
            ct2 = doc.add_table(rows=1+len(self.compliance_status), cols=3)
            ct2.style = 'Table Grid'
            for i,h in enumerate(['Framework','Status','Observações']):
                self._docx_set_bg(ct2.rows[0].cells[i], '1e3a8a')
                r = ct2.rows[0].cells[i].paragraphs[0].add_run(h)
                r.bold=True; r.font.size=Pt(9); r.font.name='Arial'
                r.font.color.rgb=RGBColor(255,255,255)
            for ri,fw in enumerate(self.compliance_status, 1):
                for ci,v in enumerate([fw['name'],fw['status'],fw['issues']]):
                    ct2.rows[ri].cells[ci].paragraphs[0].add_run(v).font.size = Pt(9)
            next_sec += 1

        # ── Conclusão ─────────────────────────────────────────────────
        self._docx_heading(doc, f'{next_sec}. CONCLUSÕES E PRÓXIMOS PASSOS')
        self._docx_body(doc,
            f"Este relatório identificou {risk['critical_count']} vulnerabilidades críticas "
            f"e {risk['high_count']} de alto risco no ambiente {self.provider_name}. "
            "É fundamental executar as ações prioritárias nos prazos estabelecidos.")
        doc.add_paragraph()
        self._docx_body(doc, "Próximos Passos Recomendados:", bold=True, color='1e3a8a')
        for step in ['Reunião de alinhamento com stakeholders (48h)',
                     'Execução das ações críticas (0-7 dias)',
                     'Implementação de controles preventivos (30-60 dias)',
                     'Re-auditoria de segurança (90 dias)',
                     'Estabelecimento de programa de segurança contínua']:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(2)
            bp.add_run(step).font.size = Pt(10)

        doc.save(str(filepath))
        report_hash = self._compute_hash(str(filepath))
        (filepath.with_suffix('.sha256')).write_text(
            f"SHA-256: {report_hash}\nArquivo: {filename}\nGerado em: {self.report_date.isoformat()}\n")
        print(f"✅ DOCX: {filepath}")
        return str(filepath)


# ══════════════════════════════════════════════════════════════════════
# FUNÇÃO PÚBLICA (API compatível com versão anterior)
# ══════════════════════════════════════════════════════════════════════
def generate_executive_report(scan_data: dict,
                               client_info: dict = None,
                               output_format: str = 'both') -> dict:
    """
    Gera relatórios PDF e/ou DOCX com gráficos de criticidade.

    Args:
        scan_data:     Dados do scan (bucket, files, severity_distribution, risk_score, provider)
        client_info:   {'name': ..., 'contact': ...}
        output_format: 'pdf' | 'docx' | 'both'

    Returns:
        dict com caminhos gerados: {'pdf': '...', 'docx': '...'}
    """
    print("\n" + "="*60)
    print("🚀 INICIANDO GERAÇÃO DE RELATÓRIO EXECUTIVO")
    print("="*60)
    try:
        gen     = ProfessionalReportGenerator(scan_data, client_info)
        results = {}

        if output_format in ('pdf', 'both'):
            try:
                results['pdf'] = gen.generate_pdf()
            except Exception as e:
                print(f"❌ Erro PDF: {e}")
                import traceback; traceback.print_exc()
                results['pdf_error'] = str(e)

        if output_format in ('docx', 'both') and DOCX_AVAILABLE:
            try:
                results['docx'] = gen.generate_docx()
            except Exception as e:
                print(f"❌ Erro DOCX: {e}")
                import traceback; traceback.print_exc()
                results['docx_error'] = str(e)

        print("\n" + "="*60)
        print("✅ GERAÇÃO CONCLUÍDA")
        print("="*60 + "\n")
        return results

    except Exception as e:
        print(f"\n❌ ERRO CRÍTICO: {e}")
        import traceback; traceback.print_exc()
        raise
